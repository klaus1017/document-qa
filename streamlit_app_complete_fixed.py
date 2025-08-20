import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime, timedelta, date
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import cx_Oracle
import os
from dotenv import load_dotenv
import warnings
import io
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Optional, List, Dict, Union, Tuple
from dataclasses import dataclass
from enum import Enum

# 设置Pandas配置以处理大数据量
pd.set_option("styler.render.max_elements", 500000)
pd.set_option('display.max_rows', 10000)
pd.set_option('display.max_columns', 50)

# 页面配置只能调用一次
if not hasattr(st, '_already_set_config'):
    try:
        st.set_page_config(
            layout="wide",
            initial_sidebar_state="expanded", 
            page_title="单耗分析系统"
        )
        st._already_set_config = True
    except:
        pass

# 加载环境变量
load_dotenv()

# 初始化session state
if 'cost_llm_messages' not in st.session_state:
    st.session_state.cost_llm_messages = []

# =============================================================================
# 数据类型和配置
# =============================================================================

class DataType(Enum):
    """数据类型枚举"""
    COST = "cost"
    OUTPUT = "output" 
    QUANTITY_TREND = "quantity_trend"
    UNIT_PRICE = "unit_price"
    COMPREHENSIVE = "comprehensive"

@dataclass
class QueryConfig:
    """查询配置类"""
    table_name: str
    date_field: str
    group_fields: List[str]
    agg_fields: Dict[str, str]
    filters: List[str] = None
    order_by: List[str] = None

class UnifiedDataLoader:
    """统一数据加载器"""
    
    def __init__(self):
        self.connection = None
        self._init_common_filters()
        self._init_query_configs()
    
    def _init_common_filters(self):
        """初始化通用过滤条件"""
        self.exclude_patterns = [
            'H%', 'P%', 'L%', 'F%', 'Y%', 'ZU%', 'CF%', 
            'W%', 'U%', 'X100%', 'XTU1%', 'XSN-AG%', 'HX%', 'T0%'
        ]
        self.include_patterns = ['H-%']
        self.common_material_filter = self._build_material_filter_clause()
    
    def _build_material_filter_clause(self) -> str:
        """构建物料过滤子句"""
        conditions = []
        
        for pattern in self.exclude_patterns:
            if pattern == 'H%':
                continue
            conditions.append(f"品名 NOT LIKE '{pattern}'")
        
        conditions.append("(品名 NOT LIKE 'H%' OR 品名 LIKE 'H-%')")
        
        return " AND " + " AND ".join(conditions)
    
    def _init_query_configs(self):
        """初始化查询配置"""
        self.query_configs = {
            DataType.COST: QueryConfig(
                table_name="VI_Wlcb_cost_ai",
                date_field="日期",
                group_fields=["部門名稱", "YYYYMM", "日期", "品名"],
                agg_fields={
                    "SUM(成本)": "TOTAL_COST",
                    "SUM(數量)": "QUAN", 
                    "AVG(單價)": "STD_COST"
                },
                order_by=["日期 DESC", "部門名稱", "品名"]
            ),
            
            DataType.OUTPUT: QueryConfig(
                table_name="YDM_COST_PRODUCT_DETAIL_TEMP",
                date_field="TRAN_DATE", 
                group_fields=["DEPT_NAME_STD", "TO_CHAR(TRAN_DATE, 'YYYYMM') as YYYYMM"],
                agg_fields={
                    "SUM(WPNL_SIZE)": "WPNL_SIZE"
                },
                filters=["trim(DEPT_CODE) in (select distinct DEPT_CODE from TB_xb_DEPT_AI)"],
                order_by=["YYYYMM DESC", "DEPT_NAME_STD"]
            ),
            
            DataType.QUANTITY_TREND: QueryConfig(
                table_name="VI_Wlcb_cost_ai",
                date_field="日期",
                group_fields=["日期"],
                agg_fields={
                    "SUM(數量)": "TOTAL_QUAN"
                },
                order_by=["日期"]
            ),
            
            DataType.UNIT_PRICE: QueryConfig(
                table_name="VI_Wlcb_cost_ai",
                date_field="日期", 
                group_fields=["日期", "品名"],
                agg_fields={
                    "單價": "UNIT_PRICE"
                },
                order_by=["品名", "日期"]
            )
        }
    
    def get_db_connection(self):
        """获取数据库连接"""
        try:
            user = os.getenv('DB_USER')
            password = os.getenv('DB_PASSWORD')
            host = os.getenv('DB_HOST')
            service = os.getenv('DB_SERVICE')
            
            if not all([user, password, host, service]):
                raise ValueError("Missing required database environment variables. Please check your .env file.")
            
            connection_string = f"{user}/{password}@{host}/{service}"
            return cx_Oracle.connect(connection_string)
        except Exception as e:
            st.error(f"数据库连接错误: {str(e)}")
            raise

    def _build_base_query(self, config: QueryConfig, start_date: date, end_date: date, 
                         additional_filters: Optional[List[str]] = None) -> Tuple[str, Dict]:
        """构建基础查询SQL"""
        
        select_fields = config.group_fields.copy()
        for agg_expr, alias in config.agg_fields.items():
            select_fields.append(f"{agg_expr} as {alias}")
        
        select_clause = "SELECT " + ",\n    ".join(select_fields)
        from_clause = f"FROM {config.table_name}"
        
        where_conditions = []
        params = {}
        
        if config.date_field == "日期":
            where_conditions.append(f"{config.date_field} >= :start_date")
            where_conditions.append(f"{config.date_field} <= :end_date")
            params['start_date'] = start_date
            params['end_date'] = end_date
        else:
            where_conditions.append(f"{config.date_field} >= TO_DATE(:start_date, 'YYYY-MM-DD')")
            where_conditions.append(f"{config.date_field} <= TO_DATE(:end_date, 'YYYY-MM-DD')")
            params['start_date'] = start_date.strftime('%Y-%m-%d')
            params['end_date'] = end_date.strftime('%Y-%m-%d')
        
        if config.table_name == "VI_Wlcb_cost_ai":
            where_conditions.extend([
                "(品名 NOT LIKE 'H%' OR 品名 LIKE 'H-%')",
                "品名 NOT LIKE 'P%'",
                "品名 NOT LIKE 'L%'", 
                "品名 NOT LIKE 'F%'",
                "品名 NOT LIKE 'Y%'",
                "品名 NOT LIKE 'ZU%'",
                "品名 NOT LIKE 'CF%'",
                "品名 NOT LIKE 'W%'",
                "品名 NOT LIKE 'U%'",
                "品名 NOT LIKE 'X100%'",
                "品名 NOT LIKE 'XTU1%'",
                "品名 NOT LIKE 'XSN-AG%'",
                "品名 NOT LIKE 'HX%'",
                "品名 NOT LIKE 'T0%'"
            ])
        
        if config.filters:
            where_conditions.extend(config.filters)
        
        if additional_filters:
            where_conditions.extend(additional_filters)
        
        where_clause = "WHERE " + " AND ".join(where_conditions)
        group_by_clause = f"GROUP BY {', '.join(config.group_fields)}"
        
        order_by_clause = ""
        if config.order_by:
            order_by_clause = f"ORDER BY {', '.join(config.order_by)}"
        
        query = f"""
        {select_clause}
        {from_clause}  
        {where_clause}
        {group_by_clause}
        {order_by_clause}
        """
        
        return query.strip(), params
    
    @st.cache_data(ttl=900)
    def load_data(_self, data_type: DataType, start_date: date, end_date: date, 
                  additional_filters: Optional[List[str]] = None,
                  inv_number: Optional[str] = None) -> pd.DataFrame:
        """统一的数据加载方法"""
        
        try:
            with _self.get_db_connection() as conn:
                if data_type == DataType.COMPREHENSIVE:
                    return _self._load_comprehensive_data(conn, start_date, end_date)
                
                if data_type not in _self.query_configs:
                    raise ValueError(f"不支持的数据类型: {data_type}")
                
                config = _self.query_configs[data_type]
                filters = additional_filters.copy() if additional_filters else []
                
                if inv_number and data_type == DataType.UNIT_PRICE:
                    filters.append("品名 = :inv_number")
                
                if data_type == DataType.OUTPUT:
                    try:
                        query, params = _self._build_base_query(config, start_date, end_date, filters)
                        if inv_number:
                            params['inv_number'] = inv_number
                        df = pd.read_sql(query, conn, params=params)
                    except Exception:
                        return _self._load_output_backup(conn, start_date, end_date)
                else:
                    query, params = _self._build_base_query(config, start_date, end_date, filters)
                    if inv_number:
                        params['inv_number'] = inv_number
                    df = pd.read_sql(query, conn, params=params)
                
                df = _self._post_process_data(df, data_type)
                return df
                
        except Exception as e:
            st.error(f"加载{data_type.value}数据时出错: {str(e)}")
            return pd.DataFrame()
    
    def _load_output_backup(self, conn, start_date: date, end_date: date) -> pd.DataFrame:
        """加载产出数据的备用方法"""
        backup_query = """
        SELECT 
            DEPT_NAME_STD,
            TO_CHAR(日期, 'YYYYMM') as YYYYMM,
            SUM(WPNL_SIZE) as WPNL_SIZE
        FROM V_WLCB_COST_PRODUCT
        WHERE 日期 >= TO_DATE(:start_date, 'YYYY-MM-DD')
          AND 日期 <= TO_DATE(:end_date, 'YYYY-MM-DD') 
          and trim(DEPT_CODE) in (select distinct DEPT_CODE from TB_xb_DEPT_AI)
        GROUP BY DEPT_NAME_STD, TO_CHAR(日期, 'YYYYMM')
        ORDER BY YYYYMM DESC, DEPT_NAME_STD
        """
        
        params = {
            'start_date': start_date.strftime('%Y-%m-%d'),
            'end_date': end_date.strftime('%Y-%m-%d')
        }
        
        return pd.read_sql(backup_query, conn, params=params)
    
    def _load_comprehensive_data(self, conn, start_date: date, end_date: date) -> pd.DataFrame:
        """加载综合数据（成本+产出）"""
        
        cost_df = self.load_data(DataType.COST, start_date, end_date)
        output_df = self.load_data(DataType.OUTPUT, start_date, end_date)
        
        if cost_df.empty or output_df.empty:
            return pd.DataFrame()
        
        merged_df = cost_df.merge(
            output_df[['DEPT_NAME_STD', 'YYYYMM', 'WPNL_SIZE']], 
            on=['DEPT_NAME_STD', 'YYYYMM'], 
            how='left'
        )
        
        merged_df['WPNL_SIZE'] = merged_df['WPNL_SIZE'].fillna(0)
        merged_df['UNIT_CONSUMPTION'] = merged_df['TOTAL_COST'] / merged_df['WPNL_SIZE']
        merged_df['UNIT_CONSUMPTION'] = merged_df['UNIT_CONSUMPTION'].replace([np.inf, -np.inf], 0).fillna(0)
        
        return merged_df
    
    def _post_process_data(self, df: pd.DataFrame, data_type: DataType) -> pd.DataFrame:
        """数据后处理"""
        if df.empty:
            return df
        
        numeric_columns = []
        
        if data_type == DataType.COST:
            numeric_columns = ['QUAN', 'STD_COST', 'TOTAL_COST']
            df['YEAR_MONTH'] = pd.to_datetime(df['YYYYMM'], format='%Y%m').dt.to_period('M')
            df['INV_PART_DESCRIPTION'] = df['品名']
            df['DEPT_CODE_STD'] = df['部門名稱']
            df['DEPT_NAME_STD'] = df['部門名稱']
            df['INV'] = df['品名']
            
            if '日期' in df.columns:
                df['DATE'] = pd.to_datetime(df['日期'])
                
        elif data_type == DataType.OUTPUT:
            numeric_columns = ['WPNL_SIZE']
            df['YEAR_MONTH'] = pd.to_datetime(df['YYYYMM'], format='%Y%m').dt.to_period('M')
            
        elif data_type == DataType.QUANTITY_TREND:
            numeric_columns = ['TOTAL_QUAN']
            df['DATE'] = pd.to_datetime(df['日期'])
            
        elif data_type == DataType.UNIT_PRICE:
            numeric_columns = ['UNIT_PRICE'] 
            df['DATE'] = pd.to_datetime(df['日期'])
            df['INV'] = df['品名']
        
        for col in numeric_columns:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
        
        return df

    @st.cache_data(ttl=900)
    def load_filter_options(_self, start_date: date, end_date: date) -> Tuple[pd.DataFrame, pd.DataFrame]:
        """加载筛选选项"""
        
        try:
            with _self.get_db_connection() as conn:
                dept_query = """
                SELECT DISTINCT 部門名稱 as DEPT_NAME
                FROM VI_Wlcb_cost_ai 
                WHERE YYYYMM >= TO_CHAR(:start_date, 'YYYYMM')
                  AND YYYYMM <= TO_CHAR(:end_date, 'YYYYMM')
                """ + _self.common_material_filter + """
                ORDER BY 部門名稱
                """
                
                dept_df = pd.read_sql(dept_query, conn, params={
                    'start_date': start_date,
                    'end_date': end_date
                })
                
                inv_query = """
                SELECT DISTINCT 品名 as INV
                FROM VI_Wlcb_cost_ai
                WHERE YYYYMM >= TO_CHAR(:start_date, 'YYYYMM')
                  AND YYYYMM <= TO_CHAR(:end_date, 'YYYYMM')
                """ + _self.common_material_filter + """
                ORDER BY 品名
                """
                
                inv_df = pd.read_sql(inv_query, conn, params={
                    'start_date': start_date,
                    'end_date': end_date
                })
                
                if not dept_df.empty:
                    dept_options = dept_df[['DEPT_NAME']].drop_duplicates().sort_values('DEPT_NAME')
                    dept_options['DEPT_CODE_STD'] = dept_options['DEPT_NAME']
                    dept_options['DEPT_NAME_STD'] = dept_options['DEPT_NAME']
                else:
                    dept_options = pd.DataFrame()
                
                if not inv_df.empty:
                    inv_options = inv_df[['INV']].drop_duplicates().sort_values('INV')
                    inv_options['INV_PART_DESCRIPTION'] = inv_options['INV']
                else:
                    inv_options = pd.DataFrame()
                    
                return dept_options, inv_options
                    
        except Exception as e:
            st.error(f"加载筛选选项时出错: {str(e)}")
            return pd.DataFrame(), pd.DataFrame()

# 全局统一数据加载器实例
unified_loader = UnifiedDataLoader()

# =============================================================================
# 分页管理器
# =============================================================================

class PageConfig:
    """页面配置类"""
    def __init__(self, name: str, title: str, data_types: List[DataType]):
        self.name = name
        self.title = title  
        self.data_types = data_types

class PageManager:
    """分页管理器"""
    
    def __init__(self):
        self.pages = {
            "cost_analysis": PageConfig(
                name="cost_analysis",
                title="成本分析",
                data_types=[DataType.COST, DataType.OUTPUT, DataType.COMPREHENSIVE],
            ),
            "quantity_trend": PageConfig(
                name="quantity_trend", 
                title="数量趋势",
                data_types=[DataType.QUANTITY_TREND, DataType.COST],
            ),
            "price_analysis": PageConfig(
                name="price_analysis",
                title="单价分析", 
                data_types=[DataType.UNIT_PRICE, DataType.COST],
            ),
            "comprehensive": PageConfig(
                name="comprehensive",
                title="综合分析",
                data_types=[DataType.COMPREHENSIVE, DataType.QUANTITY_TREND, DataType.UNIT_PRICE],
            )
        }
        
        self.current_page = None
        self.unified_loader = unified_loader
    
    def render_page_selector(self):
        """渲染页面选择器"""
        page_options = {}
        for page_id, config in self.pages.items():
            page_options[f" {config.title}"] = page_id
        
        selected_display = st.sidebar.selectbox(
            "选择分析页面",
            options=list(page_options.keys()),
            index=0
        )
        
        self.current_page = page_options[selected_display]
        return self.current_page
    
    def get_page_config(self, page_id: str = None) -> PageConfig:
        """获取页面配置"""
        page_id = page_id or self.current_page
        return self.pages.get(page_id)
    
    def load_page_data(self, page_id: str, start_date: date, end_date: date, 
                      filters: Optional[Dict] = None) -> Dict[str, pd.DataFrame]:
        """为指定页面加载所需的所有数据"""
        config = self.get_page_config(page_id)
        if not config:
            return {}
        
        data = {}
        for data_type in config.data_types:
            try:
                additional_filters = []
                inv_number = None
                
                if filters:
                    if filters.get('dept_filter') and data_type in [DataType.COST, DataType.COMPREHENSIVE]:
                        additional_filters.append(f"部門名稱 = '{filters['dept_filter']}'")
                    if filters.get('inv_filter') and data_type in [DataType.COST, DataType.UNIT_PRICE, DataType.COMPREHENSIVE]:
                        if data_type == DataType.UNIT_PRICE:
                            inv_number = filters['inv_filter']
                        else:
                            additional_filters.append(f"品名 = '{filters['inv_filter']}'")
                
                df = self.unified_loader.load_data(
                    data_type, 
                    start_date, 
                    end_date, 
                    additional_filters=additional_filters if additional_filters else None,
                    inv_number=inv_number
                )
                data[data_type.value] = df
                
            except Exception as e:
                st.error(f"加载 {data_type.value} 数据时出错: {str(e)}")
                data[data_type.value] = pd.DataFrame()
        
        return data

# 全局页面管理器实例
page_manager = PageManager()

# =============================================================================
# 向后兼容的适配器函数
# =============================================================================

def generate_word_report(user_question, llm_response, analysis_context):
    """生成Word格式的分析报告"""
    doc = Document()
    
    title = doc.add_heading('单耗分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(f"分析时间范围：{analysis_context.get('date_range', '未知')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(14)
    subtitle_format.italic = True
    
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('基本信息', level=1)
    
    basic_info_table = doc.add_table(rows=6, cols=2)
    basic_info_table.style = 'Table Grid'
    basic_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row in enumerate(basic_info_table.rows):
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(4)
    
    info_data = [
        ('生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ('分析时间范围', analysis_context.get('date_range', '未知')),
        ('数据量', f"{analysis_context.get('total_records', 0)} 条"),
        ('分析模型', analysis_context.get('llm_model', 'QWQ:32B')),
        ('报告版本', 'v3.0'),
        ('分析类型', '单耗分析')
    ]
    
    for i, (key, value) in enumerate(info_data):
        basic_info_table.cell(i, 0).text = key
        basic_info_table.cell(i, 1).text = str(value)
        basic_info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    doc.add_heading('分析问题', level=1)
    question_para = doc.add_paragraph()
    question_run = question_para.add_run(user_question)
    question_run.font.size = Pt(12)
    question_para.style = 'Quote'
    
    doc.add_heading('分析结果', level=1)
    
    response_paragraphs = llm_response.split('\n')
    
    for para in response_paragraphs:
        para = para.strip()
        if not para:
            continue
            
        if para.startswith('##'):
            heading_text = para.lstrip('#').strip()
            doc.add_heading(heading_text, level=2)
        elif para.startswith('#'):
            heading_text = para.lstrip('#').strip()
            doc.add_heading(heading_text, level=2)
        elif para.startswith('**') and para.endswith('**'):
            bold_para = doc.add_paragraph()
            bold_run = bold_para.add_run(para.strip('*'))
            bold_run.bold = True
        elif para.startswith('- ') or para.startswith('* '):
            doc.add_paragraph(para[2:], style='List Bullet')
        elif para.startswith(('1. ', '2. ', '3. ', '4. ', '5. ')):
            doc.add_paragraph(para[3:], style='List Number')
        else:
            normal_para = doc.add_paragraph(para)
            normal_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if 'data_summary' in analysis_context:
        doc.add_heading('数据概览', level=1)
        data_summary = analysis_context['data_summary']
        summary_table = doc.add_table(rows=len(data_summary), cols=2)
        summary_table.style = 'Table Grid'
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, (key, value) in enumerate(data_summary.items()):
            summary_table.cell(i, 0).text = key
            summary_table.cell(i, 1).text = str(value)
            summary_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    doc.add_heading('技术说明', level=1)
    tech_notes = [
        "本报告基于AI大模型分析生成，结合了统计学方法和机器学习技术",
        "数据来源于Oracle数据库中的成本和产出数据",
        "分析结果仅供参考，重要决策请结合专家判断",
        "建议定期更新数据并重新生成分析报告"
    ]
    
    for note in tech_notes:
        doc.add_paragraph(note, style='List Bullet')
    
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_run1 = footer_para.add_run('本报告由单耗分析系统自动生成\n')
    footer_run1.font.size = Pt(10)
    footer_run1.italic = True
    
    footer_run2 = footer_para.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
    footer_run2.font.size = Pt(8)
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def get_db_connection():
    """创建并返回Oracle数据库连接 - 统一接口"""
    return unified_loader.get_db_connection()

@st.cache_data(ttl=900)
def load_cost_data_from_db(start_date, end_date):
    """从VI_Wlcb_cost_ai表加载成本数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.COST, start_date, end_date)

@st.cache_data(ttl=900)
def load_output_data_from_db(start_date, end_date):
    """从YDM_COST_PRODUCT_DETAIL_TEMP表加载产出数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.OUTPUT, start_date, end_date)

@st.cache_data(ttl=900)
def load_data_from_db(start_date, end_date):
    """合并成本和产出数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.COMPREHENSIVE, start_date, end_date)

@st.cache_data(ttl=900)
def load_quantity_trend_data(start_date, end_date):
    """专门加载数量趋势数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.QUANTITY_TREND, start_date, end_date)

@st.cache_data(ttl=900)
def load_unit_price_data(start_date, end_date, inv_number=None):
    """专门加载单价数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.UNIT_PRICE, start_date, end_date, inv_number=inv_number)

@st.cache_data(ttl=900)
def load_filter_options(start_date, end_date):
    """加载筛选选项 - 使用统一数据加载器"""
    return unified_loader.load_filter_options(start_date, end_date)

@st.cache_data
def get_filter_options(df):
    """从DataFrame中获取筛选选项"""
    if df.empty:
        return pd.DataFrame(), pd.DataFrame()
    
    dept_options = df[['DEPT_CODE_STD', 'DEPT_NAME_STD']].drop_duplicates().sort_values('DEPT_NAME_STD')
    inv_options = df[['INV', 'INV_PART_DESCRIPTION']].drop_duplicates().sort_values('INV_PART_DESCRIPTION')
    
    return dept_options, inv_options

# =============================================================================
# 分析功能函数
# =============================================================================

def calculate_monthly_summary(df):
    """按月计算汇总数据 - 基于分开查询的逻辑"""
    if df.empty:
        return pd.DataFrame(), pd.DataFrame()
    
    df_work = df.copy()
    
    # 物料级别的汇总数据（已经包含单耗）
    monthly_material_summary = df_work.copy()
    
    # 部门级别的汇总数据
    monthly_dept_summary = df_work.groupby(['YEAR_MONTH', 'DEPT_CODE_STD', 'DEPT_NAME_STD']).agg({
        'TOTAL_COST': 'sum',     # 部门总成本
        'WPNL_SIZE': 'first',    # 部门产出（所有行都相同）
        'QUAN': 'sum',           # 部门总数量
        'INV': 'nunique'         # 部门物料种类数
    }).reset_index()
    
    # 计算部门单耗
    monthly_dept_summary['UNIT_CONSUMPTION'] = monthly_dept_summary['TOTAL_COST'] / monthly_dept_summary['WPNL_SIZE']
    monthly_dept_summary['UNIT_CONSUMPTION'] = monthly_dept_summary['UNIT_CONSUMPTION'].replace([np.inf, -np.inf], 0).fillna(0)
    
    return monthly_material_summary, monthly_dept_summary

def calculate_comprehensive_variance_analysis(df):
    """根据新的业务逻辑进行全面的变差分析"""
    if df.empty:
        return {}
    
    results = {}
    
    # 1. 单耗变差分析（与上月对比）
    if 'DATE' in df.columns:
        # 按物料、部门、日期计算日单耗
        daily_data = df.groupby(['INV', 'DEPT_NAME_STD', 'DATE']).agg({
            'TOTAL_COST': 'sum',
            'QUAN': 'sum'
        }).reset_index()
        
        # 获取部门日产出
        dept_daily_output = df.groupby(['DEPT_NAME_STD', 'DATE'])['WPNL_SIZE'].first().reset_index()
        
        # 合并计算日单耗
        daily_merged = daily_data.merge(dept_daily_output, on=['DEPT_NAME_STD', 'DATE'], how='left')
        daily_merged['日单耗'] = (daily_merged['TOTAL_COST'] / daily_merged['WPNL_SIZE']).replace([np.inf, -np.inf], 0).fillna(0)
        
        # 按月汇总计算月单耗
        daily_merged['年月'] = pd.to_datetime(daily_merged['DATE']).dt.to_period('M')
        monthly_unit_data = daily_merged.groupby(['INV', 'DEPT_NAME_STD', '年月']).agg({
            'TOTAL_COST': 'sum',
            'WPNL_SIZE': 'sum'
        }).reset_index()
        monthly_unit_data['月单耗'] = (monthly_unit_data['TOTAL_COST'] / monthly_unit_data['WPNL_SIZE']).replace([np.inf, -np.inf], 0).fillna(0)
        
        # 计算单耗变差（与上月对比）
        monthly_unit_sorted = monthly_unit_data.sort_values(['INV', 'DEPT_NAME_STD', '年月'])
        monthly_unit_sorted['上月单耗'] = monthly_unit_sorted.groupby(['INV', 'DEPT_NAME_STD'])['月单耗'].shift(1)
        monthly_unit_sorted['单耗变差'] = monthly_unit_sorted['月单耗'] - monthly_unit_sorted['上月单耗']
        monthly_unit_sorted['单耗变差百分比'] = ((monthly_unit_sorted['单耗变差'] / monthly_unit_sorted['上月单耗']) * 100).replace([np.inf, -np.inf], np.nan).round(2)
        
        # 筛选最新月份且单耗增加的记录
        latest_month = monthly_unit_sorted['年月'].max()
        unit_variance = monthly_unit_sorted[
            (monthly_unit_sorted['年月'] == latest_month) & 
            (monthly_unit_sorted['单耗变差百分比'].notna()) &
            (monthly_unit_sorted['单耗变差'] > 0)
        ].copy()
        
        # 排序：先按变差绝对值，再按变差百分比绝对值
        unit_variance['变差绝对值'] = unit_variance['单耗变差'].abs()
        unit_variance['百分比绝对值'] = unit_variance['单耗变差百分比'].abs()
        unit_variance = unit_variance.sort_values(['变差绝对值', '百分比绝对值'], ascending=[False, False])
        
        results['单耗变差'] = unit_variance.head(10)
        results['日单耗数据'] = daily_merged
    
    # 2. 耗用量变差分析（与前一天对比）
    if 'DATE' in df.columns:
        daily_qty = df.groupby(['INV', 'DATE'])['QUAN'].sum().reset_index()
        daily_qty_sorted = daily_qty.sort_values(['INV', 'DATE'])
        daily_qty_sorted['前一天数量'] = daily_qty_sorted.groupby('INV')['QUAN'].shift(1)
        daily_qty_sorted['数量变差'] = daily_qty_sorted['QUAN'] - daily_qty_sorted['前一天数量']
        daily_qty_sorted['数量变差百分比'] = ((daily_qty_sorted['数量变差'] / daily_qty_sorted['前一天数量']) * 100).replace([np.inf, -np.inf], np.nan).round(2)
        
        # 筛选最新日期且数量增加的记录
        latest_date = daily_qty_sorted['DATE'].max()
        qty_variance = daily_qty_sorted[
            (daily_qty_sorted['DATE'] == latest_date) & 
            (daily_qty_sorted['数量变差百分比'].notna()) &
            (daily_qty_sorted['数量变差'] > 0)
        ].copy()
        
        # 排序：先按变差绝对值，再按变差百分比绝对值
        qty_variance['变差绝对值'] = qty_variance['数量变差'].abs()
        qty_variance['百分比绝对值'] = qty_variance['数量变差百分比'].abs()
        qty_variance = qty_variance.sort_values(['变差绝对值', '百分比绝对值'], ascending=[False, False])
        
        results['数量变差'] = qty_variance.head(10)
        results['每日数量数据'] = daily_qty_sorted
    
    # 3. 月耗用量变差分析（与上月对比）
    df_copy = df.copy()
    df_copy['年月'] = pd.to_datetime(df_copy['YYYYMM'] + '01', format='%Y%m%d').dt.to_period('M')
    monthly_qty = df_copy.groupby(['INV', '年月'])['QUAN'].sum().reset_index()
    monthly_qty_sorted = monthly_qty.sort_values(['INV', '年月'])
    monthly_qty_sorted['上月数量'] = monthly_qty_sorted.groupby('INV')['QUAN'].shift(1)
    monthly_qty_sorted['月数量变差'] = monthly_qty_sorted['QUAN'] - monthly_qty_sorted['上月数量']
    monthly_qty_sorted['月数量变差百分比'] = ((monthly_qty_sorted['月数量变差'] / monthly_qty_sorted['上月数量']) * 100).replace([np.inf, -np.inf], np.nan).round(2)
    
    # 筛选最新月份且数量增加的记录
    latest_month = monthly_qty_sorted['年月'].max()
    monthly_qty_variance = monthly_qty_sorted[
        (monthly_qty_sorted['年月'] == latest_month) & 
        (monthly_qty_sorted['月数量变差百分比'].notna()) &
        (monthly_qty_sorted['月数量变差'] > 0)
    ].copy()
    
    # 排序：先按变差绝对值，再按变差百分比绝对值
    monthly_qty_variance['变差绝对值'] = monthly_qty_variance['月数量变差'].abs()
    monthly_qty_variance['百分比绝对值'] = monthly_qty_variance['月数量变差百分比'].abs()
    monthly_qty_variance = monthly_qty_variance.sort_values(['变差绝对值', '百分比绝对值'], ascending=[False, False])
    
    results['月数量变差'] = monthly_qty_variance.head(10)
    
    return results