import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime, timedelta, date
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import cx_Oracle
import os
from dotenv import load_dotenv
import warnings
import io
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Optional, List, Dict, Union, Tuple
from dataclasses import dataclass
from enum import Enum

# 设置Pandas配置以处理大数据量
pd.set_option("styler.render.max_elements", 500000)  # 增加Styler渲染限制
pd.set_option('display.max_rows', 10000)  # 增加显示行数限制
pd.set_option('display.max_columns', 50)  # 增加显示列数限制

# 页面配置只能调用一次，检查是否已经在主入口设置过
if not os.getenv('STREAMLIT_CONFIG_SET'):
    try:
        st.set_page_config(
            layout="wide",
            initial_sidebar_state="expanded", 
            page_title="单耗分析系统"
        )
        os.environ['STREAMLIT_CONFIG_SET'] = 'True'
    except:
        pass  # 如果已经设置过配置，忽略错误

# 加载环境变量
load_dotenv()

# =============================================================================
# 统一数据加载框架
# =============================================================================

class DataType(Enum):
    """数据类型枚举"""
    COST = "cost"
    OUTPUT = "output" 
    QUANTITY_TREND = "quantity_trend"
    UNIT_PRICE = "unit_price"
    COMPREHENSIVE = "comprehensive"

@dataclass
class QueryConfig:
    """查询配置类"""
    table_name: str
    date_field: str
    group_fields: List[str]
    agg_fields: Dict[str, str]
    filters: List[str] = None
    order_by: List[str] = None

class UnifiedDataLoader:
    """统一数据加载器"""
    
    def __init__(self):
        self.connection = None
        self._init_common_filters()
        self._init_query_configs()
    
    def _init_common_filters(self):
        """初始化通用过滤条件"""
        self.exclude_patterns = [
            'H%', 'P%', 'L%', 'F%', 'Y%', 'ZU%', 'CF%', 
            'W%', 'U%', 'X100%', 'XTU1%', 'XSN-AG%', 'HX%', 'T0%'
        ]
        self.include_patterns = ['H-%']
        
        # 构建通用WHERE子句
        self.common_material_filter = self._build_material_filter_clause()
    
    def _build_material_filter_clause(self) -> str:
        """构建物料过滤子句"""
        conditions = []
        
        # 排除条件
        for pattern in self.exclude_patterns:
            if pattern == 'H%':
                continue  # H% 需要特殊处理
            conditions.append(f"品名 NOT LIKE '{pattern}'")
        
        # H% 的特殊处理：排除H%但包含H-%
        conditions.append("(品名 NOT LIKE 'H%' OR 品名 LIKE 'H-%')")
        
        return " AND " + " AND ".join(conditions)
    
    def _init_query_configs(self):
        """初始化查询配置"""
        self.query_configs = {
            DataType.COST: QueryConfig(
                table_name="VI_Wlcb_cost_ai",
                date_field="日期",
                group_fields=["部門名稱", "YYYYMM", "日期", "品名"],
                agg_fields={
                    "SUM(成本)": "TOTAL_COST",
                    "SUM(數量)": "QUAN", 
                    "AVG(單價)": "STD_COST"
                },
                order_by=["日期 DESC", "部門名稱", "品名"]
            ),
            
            DataType.OUTPUT: QueryConfig(
                table_name="YDM_COST_PRODUCT_DETAIL_TEMP",
                date_field="TRAN_DATE", 
                group_fields=["DEPT_NAME_STD", "TO_CHAR(TRAN_DATE, 'YYYYMM') as YYYYMM"],
                agg_fields={
                    "SUM(WPNL_SIZE)": "WPNL_SIZE"
                },
                filters=["trim(DEPT_CODE) in (select distinct DEPT_CODE from TB_xb_DEPT_AI)"],
                order_by=["YYYYMM DESC", "DEPT_NAME_STD"]
            ),
            
            DataType.QUANTITY_TREND: QueryConfig(
                table_name="VI_Wlcb_cost_ai",
                date_field="日期",
                group_fields=["日期"],
                agg_fields={
                    "SUM(數量)": "TOTAL_QUAN"
                },
                order_by=["日期"]
            ),
            
            DataType.UNIT_PRICE: QueryConfig(
                table_name="VI_Wlcb_cost_ai",
                date_field="日期", 
                group_fields=["日期", "品名"],
                agg_fields={
                    "單價": "UNIT_PRICE"
                },
                order_by=["品名", "日期"]
            )
        }
    
    def get_db_connection(self):
        """获取数据库连接"""
        try:
            user = os.getenv('DB_USER')
            password = os.getenv('DB_PASSWORD')
            host = os.getenv('DB_HOST')
            service = os.getenv('DB_SERVICE')
            
            if not all([user, password, host, service]):
                raise ValueError("Missing required database environment variables. Please check your .env file.")
            
            connection_string = f"{user}/{password}@{host}/{service}"
            return cx_Oracle.connect(connection_string)
        except Exception as e:
            st.error(f"数据库连接错误: {str(e)}")
            raise

    def _build_base_query(self, config: QueryConfig, start_date: date, end_date: date, 
                         additional_filters: Optional[List[str]] = None) -> Tuple[str, Dict]:
        """构建基础查询SQL"""
        
        # SELECT子句
        select_fields = config.group_fields.copy()
        for agg_expr, alias in config.agg_fields.items():
            select_fields.append(f"{agg_expr} as {alias}")
        
        select_clause = "SELECT " + ",\n    ".join(select_fields)
        
        # FROM子句
        from_clause = f"FROM {config.table_name}"
        
        # WHERE子句
        where_conditions = []
        params = {}
        
        # 日期条件
        if config.date_field == "日期":
            where_conditions.append(f"{config.date_field} >= :start_date")
            where_conditions.append(f"{config.date_field} <= :end_date")
            params['start_date'] = start_date
            params['end_date'] = end_date
        else:  # TRAN_DATE等其他日期字段
            where_conditions.append(f"{config.date_field} >= TO_DATE(:start_date, 'YYYY-MM-DD')")
            where_conditions.append(f"{config.date_field} <= TO_DATE(:end_date, 'YYYY-MM-DD')")
            params['start_date'] = start_date.strftime('%Y-%m-%d')
            params['end_date'] = end_date.strftime('%Y-%m-%d')
        
        # 物料过滤条件（仅对包含品名字段的表应用）
        if config.table_name == "VI_Wlcb_cost_ai":
            where_conditions.extend([
                "(品名 NOT LIKE 'H%' OR 品名 LIKE 'H-%')",
                "品名 NOT LIKE 'P%'",
                "品名 NOT LIKE 'L%'", 
                "品名 NOT LIKE 'F%'",
                "品名 NOT LIKE 'Y%'",
                "品名 NOT LIKE 'ZU%'",
                "品名 NOT LIKE 'CF%'",
                "品名 NOT LIKE 'W%'",
                "品名 NOT LIKE 'U%'",
                "品名 NOT LIKE 'X100%'",
                "品名 NOT LIKE 'XTU1%'",
                "品名 NOT LIKE 'XSN-AG%'",
                "品名 NOT LIKE 'HX%'",
                "品名 NOT LIKE 'T0%'"
            ])
        
        # 配置中的额外过滤条件
        if config.filters:
            where_conditions.extend(config.filters)
        
        # 动态添加的过滤条件
        if additional_filters:
            where_conditions.extend(additional_filters)
        
        where_clause = "WHERE " + " AND ".join(where_conditions)
        
        # GROUP BY子句
        group_by_clause = f"GROUP BY {', '.join(config.group_fields)}"
        
        # ORDER BY子句
        order_by_clause = ""
        if config.order_by:
            order_by_clause = f"ORDER BY {', '.join(config.order_by)}"
        
        # 组装完整SQL
        query = f"""
        {select_clause}
        {from_clause}  
        {where_clause}
        {group_by_clause}
        {order_by_clause}
        """
        
        return query.strip(), params
    
    @st.cache_data(ttl=900)
    def load_data(_self, data_type: DataType, start_date: date, end_date: date, 
                  additional_filters: Optional[List[str]] = None,
                  inv_number: Optional[str] = None) -> pd.DataFrame:
        """
        统一的数据加载方法
        
        Args:
            data_type: 数据类型
            start_date: 开始日期
            end_date: 结束日期  
            additional_filters: 额外的过滤条件
            inv_number: 指定的料号（用于单价查询）
        
        Returns:
            pd.DataFrame: 查询结果
        """
        
        try:
            with _self.get_db_connection() as conn:
                # 处理特殊情况：综合数据需要合并多个查询
                if data_type == DataType.COMPREHENSIVE:
                    return _self._load_comprehensive_data(conn, start_date, end_date)
                
                # 获取查询配置
                if data_type not in _self.query_configs:
                    raise ValueError(f"不支持的数据类型: {data_type}")
                
                config = _self.query_configs[data_type]
                
                # 处理特殊过滤条件
                filters = additional_filters.copy() if additional_filters else []
                
                # 处理料号过滤（用于单价查询）
                if inv_number and data_type == DataType.UNIT_PRICE:
                    filters.append("品名 = :inv_number")
                
                # 处理输出数据的备用查询
                if data_type == DataType.OUTPUT:
                    try:
                        # 首先尝试主表
                        query, params = _self._build_base_query(config, start_date, end_date, filters)
                        if inv_number:
                            params['inv_number'] = inv_number
                        df = pd.read_sql(query, conn, params=params)
                    except Exception:
                        # 如果主表查询失败，尝试备用表
                        return _self._load_output_backup(conn, start_date, end_date)
                else:
                    # 构建并执行查询
                    query, params = _self._build_base_query(config, start_date, end_date, filters)
                    if inv_number:
                        params['inv_number'] = inv_number
                    df = pd.read_sql(query, conn, params=params)
                
                # 数据后处理
                df = _self._post_process_data(df, data_type)
                
                return df
                
        except Exception as e:
            st.error(f"加载{data_type.value}数据时出错: {str(e)}")
            return pd.DataFrame()
    
    def _load_output_backup(self, conn, start_date: date, end_date: date) -> pd.DataFrame:
        """加载产出数据的备用方法"""
        backup_query = """
        SELECT 
            DEPT_NAME_STD,
            TO_CHAR(日期, 'YYYYMM') as YYYYMM,
            SUM(WPNL_SIZE) as WPNL_SIZE
        FROM V_WLCB_COST_PRODUCT
        WHERE 日期 >= TO_DATE(:start_date, 'YYYY-MM-DD')
          AND 日期 <= TO_DATE(:end_date, 'YYYY-MM-DD') 
          and trim(DEPT_CODE) in (select distinct DEPT_CODE from TB_xb_DEPT_AI)
        GROUP BY DEPT_NAME_STD, TO_CHAR(日期, 'YYYYMM')
        ORDER BY YYYYMM DESC, DEPT_NAME_STD
        """
        
        params = {
            'start_date': start_date.strftime('%Y-%m-%d'),
            'end_date': end_date.strftime('%Y-%m-%d')
        }
        
        return pd.read_sql(backup_query, conn, params=params)
    
    def _load_comprehensive_data(self, conn, start_date: date, end_date: date) -> pd.DataFrame:
        """加载综合数据（成本+产出）"""
        
        # 加载成本数据
        cost_df = self.load_data(DataType.COST, start_date, end_date)
        
        # 加载产出数据  
        output_df = self.load_data(DataType.OUTPUT, start_date, end_date)
        
        if cost_df.empty:
            return pd.DataFrame()
        
        if output_df.empty:
            return pd.DataFrame()
        
        # 合并成本和产出数据
        merged_df = cost_df.merge(
            output_df[['DEPT_NAME_STD', 'YYYYMM', 'WPNL_SIZE']], 
            on=['DEPT_NAME_STD', 'YYYYMM'], 
            how='left'
        )
        
        # 填充缺失的产出数据并计算单耗
        merged_df['WPNL_SIZE'] = merged_df['WPNL_SIZE'].fillna(0)
        merged_df['UNIT_CONSUMPTION'] = merged_df['TOTAL_COST'] / merged_df['WPNL_SIZE']
        merged_df['UNIT_CONSUMPTION'] = merged_df['UNIT_CONSUMPTION'].replace([np.inf, -np.inf], 0).fillna(0)
        
        return merged_df
    
    def _post_process_data(self, df: pd.DataFrame, data_type: DataType) -> pd.DataFrame:
        """数据后处理"""
        if df.empty:
            return df
        
        # 通用数据类型转换
        numeric_columns = []
        
        if data_type == DataType.COST:
            # 成本数据处理
            numeric_columns = ['QUAN', 'STD_COST', 'TOTAL_COST']
            
            # 添加计算字段
            df['YEAR_MONTH'] = pd.to_datetime(df['YYYYMM'], format='%Y%m').dt.to_period('M')
            df['INV_PART_DESCRIPTION'] = df['品名']
            df['DEPT_CODE_STD'] = df['部門名稱']
            df['DEPT_NAME_STD'] = df['部門名稱']
            df['INV'] = df['品名']
            
            # 处理日期字段
            if '日期' in df.columns:
                df['DATE'] = pd.to_datetime(df['日期'])
                
        elif data_type == DataType.OUTPUT:
            # 产出数据处理
            numeric_columns = ['WPNL_SIZE']
            df['YEAR_MONTH'] = pd.to_datetime(df['YYYYMM'], format='%Y%m').dt.to_period('M')
            
        elif data_type == DataType.QUANTITY_TREND:
            # 数量趋势数据处理  
            numeric_columns = ['TOTAL_QUAN']
            df['DATE'] = pd.to_datetime(df['日期'])
            
        elif data_type == DataType.UNIT_PRICE:
            # 单价数据处理
            numeric_columns = ['UNIT_PRICE'] 
            df['DATE'] = pd.to_datetime(df['日期'])
            df['INV'] = df['品名']
        
        # 数值类型转换
        for col in numeric_columns:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
        
        return df

    @st.cache_data(ttl=900)
    def load_filter_options(_self, start_date: date, end_date: date) -> Tuple[pd.DataFrame, pd.DataFrame]:
        """加载筛选选项"""
        
        try:
            with _self.get_db_connection() as conn:
                # 加载部门选项
                dept_query = """
                SELECT DISTINCT 部門名稱 as DEPT_NAME
                FROM VI_Wlcb_cost_ai 
                WHERE YYYYMM >= TO_CHAR(:start_date, 'YYYYMM')
                  AND YYYYMM <= TO_CHAR(:end_date, 'YYYYMM')
                """ + _self.common_material_filter + """
                ORDER BY 部門名稱
                """
                
                dept_df = pd.read_sql(dept_query, conn, params={
                    'start_date': start_date,
                    'end_date': end_date
                })
                
                # 加载物料选项
                inv_query = """
                SELECT DISTINCT 品名 as INV
                FROM VI_Wlcb_cost_ai
                WHERE YYYYMM >= TO_CHAR(:start_date, 'YYYYMM')
                  AND YYYYMM <= TO_CHAR(:end_date, 'YYYYMM')
                """ + _self.common_material_filter + """
                ORDER BY 品名
                """
                
                inv_df = pd.read_sql(inv_query, conn, params={
                    'start_date': start_date,
                    'end_date': end_date
                })
                
                # 处理部门选项
                if not dept_df.empty:
                    dept_options = dept_df[['DEPT_NAME']].drop_duplicates().sort_values('DEPT_NAME')
                    dept_options['DEPT_CODE_STD'] = dept_options['DEPT_NAME']
                    dept_options['DEPT_NAME_STD'] = dept_options['DEPT_NAME']
                else:
                    dept_options = pd.DataFrame()
                
                # 处理物料选项
                if not inv_df.empty:
                    inv_options = inv_df[['INV']].drop_duplicates().sort_values('INV')
                    inv_options['INV_PART_DESCRIPTION'] = inv_options['INV']
                else:
                    inv_options = pd.DataFrame()
                    
                return dept_options, inv_options
                    
        except Exception as e:
            st.error(f"加载筛选选项时出错: {str(e)}")
            return pd.DataFrame(), pd.DataFrame()

# 全局统一数据加载器实例
unified_loader = UnifiedDataLoader()

# =============================================================================
# 向后兼容的适配器函数 (替换原有的数据加载函数)
# =============================================================================

def generate_word_report(user_question, llm_response, analysis_context):
    """生成Word格式的分析报告"""
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading('单耗分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加副标题
    subtitle = doc.add_paragraph(f"分析时间范围：{analysis_context.get('date_range', '未知')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(14)
    subtitle_format.italic = True
    
    # 添加分隔线
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加基本信息
    doc.add_heading('基本信息', level=1)
    
    # 创建基本信息表格
    basic_info_table = doc.add_table(rows=6, cols=2)
    basic_info_table.style = 'Table Grid'
    basic_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表格列宽
    for i, row in enumerate(basic_info_table.rows):
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(4)
    
    # 填充基本信息
    info_data = [
        ('生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ('分析时间范围', analysis_context.get('date_range', '未知')),
        ('数据量', f"{analysis_context.get('total_records', 0)} 条"),
        ('分析模型', analysis_context.get('llm_model', 'QWQ:32B')),
        ('报告版本', 'v3.0'),
        ('分析类型', '单耗分析')
    ]
    
    for i, (key, value) in enumerate(info_data):
        basic_info_table.cell(i, 0).text = key
        basic_info_table.cell(i, 1).text = str(value)
        # 设置左列为粗体
        basic_info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    # 添加分析问题
    doc.add_heading('分析问题', level=1)
    question_para = doc.add_paragraph()
    question_run = question_para.add_run(user_question)
    question_run.font.size = Pt(12)
    question_para.style = 'Quote'
    
    # 添加分析结果
    doc.add_heading('分析结果', level=1)
    
    # 处理LLM回复，按段落分割并改进格式化
    response_paragraphs = llm_response.split('\n')
    
    for para in response_paragraphs:
        para = para.strip()
        if not para:
            continue
            
        # 检查是否是Markdown标题
        if para.startswith('##'):
            # 二级标题
            heading_text = para.lstrip('#').strip()
            doc.add_heading(heading_text, level=2)
        elif para.startswith('#'):
            # 一级标题
            heading_text = para.lstrip('#').strip()
            doc.add_heading(heading_text, level=2)
        elif para.startswith('**') and para.endswith('**'):
            # 粗体段落
            bold_para = doc.add_paragraph()
            bold_run = bold_para.add_run(para.strip('*'))
            bold_run.bold = True
        elif para.startswith('- ') or para.startswith('* '):
            # 列表项
            list_para = doc.add_paragraph(para[2:], style='List Bullet')
        elif para.startswith(('1. ', '2. ', '3. ', '4. ', '5. ')):
            # 数字列表项
            list_para = doc.add_paragraph(para[3:], style='List Number')
        else:
            # 普通段落
            normal_para = doc.add_paragraph(para)
            normal_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # 添加数据概览
    if 'data_summary' in analysis_context:
        doc.add_heading('数据概览', level=1)
        data_summary = analysis_context['data_summary']
        summary_table = doc.add_table(rows=len(data_summary), cols=2)
        summary_table.style = 'Table Grid'
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表格样式
        for i, (key, value) in enumerate(data_summary.items()):
            summary_table.cell(i, 0).text = key
            summary_table.cell(i, 1).text = str(value)
            # 设置左列为粗体
            summary_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    # 添加技术说明
    doc.add_heading('技术说明', level=1)
    tech_notes = [
        "本报告基于AI大模型分析生成，结合了统计学方法和机器学习技术",
        "数据来源于Oracle数据库中的成本和产出数据",
        "分析结果仅供参考，重要决策请结合专家判断",
        "建议定期更新数据并重新生成分析报告"
    ]
    
    for note in tech_notes:
        note_para = doc.add_paragraph(note, style='List Bullet')
    
    # 添加页脚
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_run1 = footer_para.add_run('本报告由单耗分析系统自动生成\n')
    footer_run1.font.size = Pt(10)
    footer_run1.italic = True
    
    footer_run2 = footer_para.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
    footer_run2.font.size = Pt(8)
    footer_run2.font.color.rgb = None  # 灰色
    
    # 将文档保存到内存
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def get_db_connection():
    """创建并返回Oracle数据库连接 - 统一接口"""
    return unified_loader.get_db_connection()

@st.cache_data(ttl=900)  # 15分钟缓存
def load_cost_data_from_db(start_date, end_date):
    """从VI_Wlcb_cost_ai表加载成本数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.COST, start_date, end_date)

@st.cache_data(ttl=900)  # 15分钟缓存  
def load_output_data_from_db(start_date, end_date):
    """从YDM_COST_PRODUCT_DETAIL_TEMP表加载产出数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.OUTPUT, start_date, end_date)

@st.cache_data(ttl=900)
def load_data_from_db(start_date, end_date):
    """合并成本和产出数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.COMPREHENSIVE, start_date, end_date)

@st.cache_data(ttl=900)
def load_quantity_trend_data(start_date, end_date):
    """专门加载数量趋势数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.QUANTITY_TREND, start_date, end_date)

@st.cache_data(ttl=900)
def load_unit_price_data(start_date, end_date, inv_number=None):
    """专门加载单价数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.UNIT_PRICE, start_date, end_date, inv_number=inv_number)

@st.cache_data(ttl=900)
def load_filter_options(start_date, end_date):
    """加载筛选选项 - 使用统一数据加载器"""
    return unified_loader.load_filter_options(start_date, end_date)