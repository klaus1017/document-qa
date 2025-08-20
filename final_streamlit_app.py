import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime, timedelta, date
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import cx_Oracle
import os
from dotenv import load_dotenv
import warnings
import io
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Optional, List, Dict, Union, Tuple
from dataclasses import dataclass
from enum import Enum

# 设置Pandas配置以处理大数据量
pd.set_option("styler.render.max_elements", 500000)
pd.set_option('display.max_rows', 10000)
pd.set_option('display.max_columns', 50)

# 页面配置只能调用一次
if not hasattr(st, '_already_set_config'):
    try:
        st.set_page_config(
            layout="wide",
            initial_sidebar_state="expanded", 
            page_title="单耗分析系统"
        )
        st._already_set_config = True
    except:
        pass

# 加载环境变量
load_dotenv()

# 初始化session state
if 'cost_llm_messages' not in st.session_state:
    st.session_state.cost_llm_messages = []

# =============================================================================
# 数据类型和配置
# =============================================================================

class DataType(Enum):
    """数据类型枚举"""
    COST = "cost"
    OUTPUT = "output" 
    QUANTITY_TREND = "quantity_trend"
    UNIT_PRICE = "unit_price"
    COMPREHENSIVE = "comprehensive"

@dataclass
class QueryConfig:
    """查询配置类"""
    table_name: str
    date_field: str
    group_fields: List[str]
    agg_fields: Dict[str, str]
    filters: List[str] = None
    order_by: List[str] = None

class UnifiedDataLoader:
    """统一数据加载器"""
    
    def __init__(self):
        self.connection = None
        self._init_common_filters()
        self._init_query_configs()
    
    def _init_common_filters(self):
        """初始化通用过滤条件"""
        self.exclude_patterns = [
            'H%', 'P%', 'L%', 'F%', 'Y%', 'ZU%', 'CF%', 
            'W%', 'U%', 'X100%', 'XTU1%', 'XSN-AG%', 'HX%', 'T0%'
        ]
        self.include_patterns = ['H-%']
        self.common_material_filter = self._build_material_filter_clause()
    
    def _build_material_filter_clause(self) -> str:
        """构建物料过滤子句"""
        conditions = []
        
        for pattern in self.exclude_patterns:
            if pattern == 'H%':
                continue
            conditions.append(f"品名 NOT LIKE '{pattern}'")
        
        conditions.append("(品名 NOT LIKE 'H%' OR 品名 LIKE 'H-%')")
        
        return " AND " + " AND ".join(conditions)
    
    def _init_query_configs(self):
        """初始化查询配置"""
        self.query_configs = {
            DataType.COST: QueryConfig(
                table_name="VI_Wlcb_cost_ai",
                date_field="日期",
                group_fields=["部門名稱", "YYYYMM", "日期", "品名"],
                agg_fields={
                    "SUM(成本)": "TOTAL_COST",
                    "SUM(數量)": "QUAN", 
                    "AVG(單價)": "STD_COST"
                },
                order_by=["日期 DESC", "部門名稱", "品名"]
            ),
            
            DataType.OUTPUT: QueryConfig(
                table_name="YDM_COST_PRODUCT_DETAIL_TEMP",
                date_field="TRAN_DATE", 
                group_fields=["DEPT_NAME_STD", "TO_CHAR(TRAN_DATE, 'YYYYMM') as YYYYMM"],
                agg_fields={
                    "SUM(WPNL_SIZE)": "WPNL_SIZE"
                },
                filters=["trim(DEPT_CODE) in (select distinct DEPT_CODE from TB_xb_DEPT_AI)"],
                order_by=["YYYYMM DESC", "DEPT_NAME_STD"]
            ),
            
            DataType.QUANTITY_TREND: QueryConfig(
                table_name="VI_Wlcb_cost_ai",
                date_field="日期",
                group_fields=["日期"],
                agg_fields={
                    "SUM(數量)": "TOTAL_QUAN"
                },
                order_by=["日期"]
            ),
            
            DataType.UNIT_PRICE: QueryConfig(
                table_name="VI_Wlcb_cost_ai",
                date_field="日期", 
                group_fields=["日期", "品名"],
                agg_fields={
                    "單價": "UNIT_PRICE"
                },
                order_by=["品名", "日期"]
            )
        }
    
    def get_db_connection(self):
        """获取数据库连接"""
        try:
            user = os.getenv('DB_USER')
            password = os.getenv('DB_PASSWORD')
            host = os.getenv('DB_HOST')
            service = os.getenv('DB_SERVICE')
            
            if not all([user, password, host, service]):
                raise ValueError("Missing required database environment variables. Please check your .env file.")
            
            connection_string = f"{user}/{password}@{host}/{service}"
            return cx_Oracle.connect(connection_string)
        except Exception as e:
            st.error(f"数据库连接错误: {str(e)}")
            raise

    def _build_base_query(self, config: QueryConfig, start_date: date, end_date: date, 
                         additional_filters: Optional[List[str]] = None) -> Tuple[str, Dict]:
        """构建基础查询SQL"""
        
        select_fields = config.group_fields.copy()
        for agg_expr, alias in config.agg_fields.items():
            select_fields.append(f"{agg_expr} as {alias}")
        
        select_clause = "SELECT " + ",\n    ".join(select_fields)
        from_clause = f"FROM {config.table_name}"
        
        where_conditions = []
        params = {}
        
        if config.date_field == "日期":
            where_conditions.append(f"{config.date_field} >= :start_date")
            where_conditions.append(f"{config.date_field} <= :end_date")
            params['start_date'] = start_date
            params['end_date'] = end_date
        else:
            where_conditions.append(f"{config.date_field} >= TO_DATE(:start_date, 'YYYY-MM-DD')")
            where_conditions.append(f"{config.date_field} <= TO_DATE(:end_date, 'YYYY-MM-DD')")
            params['start_date'] = start_date.strftime('%Y-%m-%d')
            params['end_date'] = end_date.strftime('%Y-%m-%d')
        
        if config.table_name == "VI_Wlcb_cost_ai":
            where_conditions.extend([
                "(品名 NOT LIKE 'H%' OR 品名 LIKE 'H-%')",
                "品名 NOT LIKE 'P%'",
                "品名 NOT LIKE 'L%'", 
                "品名 NOT LIKE 'F%'",
                "品名 NOT LIKE 'Y%'",
                "品名 NOT LIKE 'ZU%'",
                "品名 NOT LIKE 'CF%'",
                "品名 NOT LIKE 'W%'",
                "品名 NOT LIKE 'U%'",
                "品名 NOT LIKE 'X100%'",
                "品名 NOT LIKE 'XTU1%'",
                "品名 NOT LIKE 'XSN-AG%'",
                "品名 NOT LIKE 'HX%'",
                "品名 NOT LIKE 'T0%'"
            ])
        
        if config.filters:
            where_conditions.extend(config.filters)
        
        if additional_filters:
            where_conditions.extend(additional_filters)
        
        where_clause = "WHERE " + " AND ".join(where_conditions)
        group_by_clause = f"GROUP BY {', '.join(config.group_fields)}"
        
        order_by_clause = ""
        if config.order_by:
            order_by_clause = f"ORDER BY {', '.join(config.order_by)}"
        
        query = f"""
        {select_clause}
        {from_clause}  
        {where_clause}
        {group_by_clause}
        {order_by_clause}
        """
        
        return query.strip(), params
    
    @st.cache_data(ttl=900)
    def load_data(_self, data_type: DataType, start_date: date, end_date: date, 
                  additional_filters: Optional[List[str]] = None,
                  inv_number: Optional[str] = None) -> pd.DataFrame:
        """统一的数据加载方法"""
        
        try:
            with _self.get_db_connection() as conn:
                if data_type == DataType.COMPREHENSIVE:
                    return _self._load_comprehensive_data(conn, start_date, end_date)
                
                if data_type not in _self.query_configs:
                    raise ValueError(f"不支持的数据类型: {data_type}")
                
                config = _self.query_configs[data_type]
                filters = additional_filters.copy() if additional_filters else []
                
                if inv_number and data_type == DataType.UNIT_PRICE:
                    filters.append("品名 = :inv_number")
                
                if data_type == DataType.OUTPUT:
                    try:
                        query, params = _self._build_base_query(config, start_date, end_date, filters)
                        if inv_number:
                            params['inv_number'] = inv_number
                        df = pd.read_sql(query, conn, params=params)
                    except Exception:
                        return _self._load_output_backup(conn, start_date, end_date)
                else:
                    query, params = _self._build_base_query(config, start_date, end_date, filters)
                    if inv_number:
                        params['inv_number'] = inv_number
                    df = pd.read_sql(query, conn, params=params)
                
                df = _self._post_process_data(df, data_type)
                return df
                
        except Exception as e:
            st.error(f"加载{data_type.value}数据时出错: {str(e)}")
            return pd.DataFrame()
    
    def _load_output_backup(self, conn, start_date: date, end_date: date) -> pd.DataFrame:
        """加载产出数据的备用方法"""
        backup_query = """
        SELECT 
            DEPT_NAME_STD,
            TO_CHAR(日期, 'YYYYMM') as YYYYMM,
            SUM(WPNL_SIZE) as WPNL_SIZE
        FROM V_WLCB_COST_PRODUCT
        WHERE 日期 >= TO_DATE(:start_date, 'YYYY-MM-DD')
          AND 日期 <= TO_DATE(:end_date, 'YYYY-MM-DD') 
          and trim(DEPT_CODE) in (select distinct DEPT_CODE from TB_xb_DEPT_AI)
        GROUP BY DEPT_NAME_STD, TO_CHAR(日期, 'YYYYMM')
        ORDER BY YYYYMM DESC, DEPT_NAME_STD
        """
        
        params = {
            'start_date': start_date.strftime('%Y-%m-%d'),
            'end_date': end_date.strftime('%Y-%m-%d')
        }
        
        return pd.read_sql(backup_query, conn, params=params)
    
    def _load_comprehensive_data(self, conn, start_date: date, end_date: date) -> pd.DataFrame:
        """加载综合数据（成本+产出）"""
        
        cost_df = self.load_data(DataType.COST, start_date, end_date)
        output_df = self.load_data(DataType.OUTPUT, start_date, end_date)
        
        if cost_df.empty or output_df.empty:
            return pd.DataFrame()
        
        merged_df = cost_df.merge(
            output_df[['DEPT_NAME_STD', 'YYYYMM', 'WPNL_SIZE']], 
            on=['DEPT_NAME_STD', 'YYYYMM'], 
            how='left'
        )
        
        merged_df['WPNL_SIZE'] = merged_df['WPNL_SIZE'].fillna(0)
        merged_df['UNIT_CONSUMPTION'] = merged_df['TOTAL_COST'] / merged_df['WPNL_SIZE']
        merged_df['UNIT_CONSUMPTION'] = merged_df['UNIT_CONSUMPTION'].replace([np.inf, -np.inf], 0).fillna(0)
        
        return merged_df
    
    def _post_process_data(self, df: pd.DataFrame, data_type: DataType) -> pd.DataFrame:
        """数据后处理"""
        if df.empty:
            return df
        
        numeric_columns = []
        
        if data_type == DataType.COST:
            numeric_columns = ['QUAN', 'STD_COST', 'TOTAL_COST']
            df['YEAR_MONTH'] = pd.to_datetime(df['YYYYMM'], format='%Y%m').dt.to_period('M')
            df['INV_PART_DESCRIPTION'] = df['品名']
            df['DEPT_CODE_STD'] = df['部門名稱']
            df['DEPT_NAME_STD'] = df['部門名稱']
            df['INV'] = df['品名']
            
            if '日期' in df.columns:
                df['DATE'] = pd.to_datetime(df['日期'])
                
        elif data_type == DataType.OUTPUT:
            numeric_columns = ['WPNL_SIZE']
            df['YEAR_MONTH'] = pd.to_datetime(df['YYYYMM'], format='%Y%m').dt.to_period('M')
            
        elif data_type == DataType.QUANTITY_TREND:
            numeric_columns = ['TOTAL_QUAN']
            df['DATE'] = pd.to_datetime(df['日期'])
            
        elif data_type == DataType.UNIT_PRICE:
            numeric_columns = ['UNIT_PRICE'] 
            df['DATE'] = pd.to_datetime(df['日期'])
            df['INV'] = df['品名']
        
        for col in numeric_columns:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
        
        return df

    @st.cache_data(ttl=900)
    def load_filter_options(_self, start_date: date, end_date: date) -> Tuple[pd.DataFrame, pd.DataFrame]:
        """加载筛选选项"""
        
        try:
            with _self.get_db_connection() as conn:
                dept_query = """
                SELECT DISTINCT 部門名稱 as DEPT_NAME
                FROM VI_Wlcb_cost_ai 
                WHERE YYYYMM >= TO_CHAR(:start_date, 'YYYYMM')
                  AND YYYYMM <= TO_CHAR(:end_date, 'YYYYMM')
                """ + _self.common_material_filter + """
                ORDER BY 部門名稱
                """
                
                dept_df = pd.read_sql(dept_query, conn, params={
                    'start_date': start_date,
                    'end_date': end_date
                })
                
                inv_query = """
                SELECT DISTINCT 品名 as INV
                FROM VI_Wlcb_cost_ai
                WHERE YYYYMM >= TO_CHAR(:start_date, 'YYYYMM')
                  AND YYYYMM <= TO_CHAR(:end_date, 'YYYYMM')
                """ + _self.common_material_filter + """
                ORDER BY 品名
                """
                
                inv_df = pd.read_sql(inv_query, conn, params={
                    'start_date': start_date,
                    'end_date': end_date
                })
                
                if not dept_df.empty:
                    dept_options = dept_df[['DEPT_NAME']].drop_duplicates().sort_values('DEPT_NAME')
                    dept_options['DEPT_CODE_STD'] = dept_options['DEPT_NAME']
                    dept_options['DEPT_NAME_STD'] = dept_options['DEPT_NAME']
                else:
                    dept_options = pd.DataFrame()
                
                if not inv_df.empty:
                    inv_options = inv_df[['INV']].drop_duplicates().sort_values('INV')
                    inv_options['INV_PART_DESCRIPTION'] = inv_options['INV']
                else:
                    inv_options = pd.DataFrame()
                    
                return dept_options, inv_options
                    
        except Exception as e:
            st.error(f"加载筛选选项时出错: {str(e)}")
            return pd.DataFrame(), pd.DataFrame()

# 全局统一数据加载器实例
unified_loader = UnifiedDataLoader()

# =============================================================================
# 向后兼容的适配器函数
# =============================================================================

def generate_word_report(user_question, llm_response, analysis_context):
    """生成Word格式的分析报告"""
    doc = Document()
    
    title = doc.add_heading('单耗分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(f"分析时间范围：{analysis_context.get('date_range', '未知')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(14)
    subtitle_format.italic = True
    
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('基本信息', level=1)
    
    basic_info_table = doc.add_table(rows=6, cols=2)
    basic_info_table.style = 'Table Grid'
    basic_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row in enumerate(basic_info_table.rows):
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(4)
    
    info_data = [
        ('生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ('分析时间范围', analysis_context.get('date_range', '未知')),
        ('数据量', f"{analysis_context.get('total_records', 0)} 条"),
        ('分析模型', analysis_context.get('llm_model', 'QWQ:32B')),
        ('报告版本', 'v3.0'),
        ('分析类型', '单耗分析')
    ]
    
    for i, (key, value) in enumerate(info_data):
        basic_info_table.cell(i, 0).text = key
        basic_info_table.cell(i, 1).text = str(value)
        basic_info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    doc.add_heading('分析问题', level=1)
    question_para = doc.add_paragraph()
    question_run = question_para.add_run(user_question)
    question_run.font.size = Pt(12)
    question_para.style = 'Quote'
    
    doc.add_heading('分析结果', level=1)
    
    response_paragraphs = llm_response.split('\n')
    
    for para in response_paragraphs:
        para = para.strip()
        if not para:
            continue
            
        if para.startswith('##'):
            heading_text = para.lstrip('#').strip()
            doc.add_heading(heading_text, level=2)
        elif para.startswith('#'):
            heading_text = para.lstrip('#').strip()
            doc.add_heading(heading_text, level=2)
        elif para.startswith('**') and para.endswith('**'):
            bold_para = doc.add_paragraph()
            bold_run = bold_para.add_run(para.strip('*'))
            bold_run.bold = True
        elif para.startswith('- ') or para.startswith('* '):
            doc.add_paragraph(para[2:], style='List Bullet')
        elif para.startswith(('1. ', '2. ', '3. ', '4. ', '5. ')):
            doc.add_paragraph(para[3:], style='List Number')
        else:
            normal_para = doc.add_paragraph(para)
            normal_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if 'data_summary' in analysis_context:
        doc.add_heading('数据概览', level=1)
        data_summary = analysis_context['data_summary']
        summary_table = doc.add_table(rows=len(data_summary), cols=2)
        summary_table.style = 'Table Grid'
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, (key, value) in enumerate(data_summary.items()):
            summary_table.cell(i, 0).text = key
            summary_table.cell(i, 1).text = str(value)
            summary_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    doc.add_heading('技术说明', level=1)
    tech_notes = [
        "本报告基于AI大模型分析生成，结合了统计学方法和机器学习技术",
        "数据来源于Oracle数据库中的成本和产出数据",
        "分析结果仅供参考，重要决策请结合专家判断",
        "建议定期更新数据并重新生成分析报告"
    ]
    
    for note in tech_notes:
        doc.add_paragraph(note, style='List Bullet')
    
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_run1 = footer_para.add_run('本报告由单耗分析系统自动生成\n')
    footer_run1.font.size = Pt(10)
    footer_run1.italic = True
    
    footer_run2 = footer_para.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
    footer_run2.font.size = Pt(8)
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def get_db_connection():
    """创建并返回Oracle数据库连接 - 统一接口"""
    return unified_loader.get_db_connection()

@st.cache_data(ttl=900)
def load_cost_data_from_db(start_date, end_date):
    """从VI_Wlcb_cost_ai表加载成本数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.COST, start_date, end_date)

@st.cache_data(ttl=900)
def load_output_data_from_db(start_date, end_date):
    """从YDM_COST_PRODUCT_DETAIL_TEMP表加载产出数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.OUTPUT, start_date, end_date)

@st.cache_data(ttl=900)
def load_data_from_db(start_date, end_date):
    """合并成本和产出数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.COMPREHENSIVE, start_date, end_date)

@st.cache_data(ttl=900)
def load_quantity_trend_data(start_date, end_date):
    """专门加载数量趋势数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.QUANTITY_TREND, start_date, end_date)

@st.cache_data(ttl=900)
def load_unit_price_data(start_date, end_date, inv_number=None):
    """专门加载单价数据 - 使用统一数据加载器"""
    return unified_loader.load_data(DataType.UNIT_PRICE, start_date, end_date, inv_number=inv_number)

@st.cache_data(ttl=900)
def load_filter_options(start_date, end_date):
    """加载筛选选项 - 使用统一数据加载器"""
    return unified_loader.load_filter_options(start_date, end_date)

# =============================================================================
# 分页管理器
# =============================================================================

class PageConfig:
    """页面配置类"""
    def __init__(self, name: str, title: str, data_types: List[DataType]):
        self.name = name
        self.title = title  
        self.data_types = data_types

class PageManager:
    """分页管理器"""
    
    def __init__(self):
        self.pages = {
            "cost_analysis": PageConfig(
                name="cost_analysis",
                title="成本分析",
                data_types=[DataType.COST, DataType.OUTPUT, DataType.COMPREHENSIVE],
            ),
            "quantity_trend": PageConfig(
                name="quantity_trend", 
                title="数量趋势",
                data_types=[DataType.QUANTITY_TREND, DataType.COST],
            ),
            "price_analysis": PageConfig(
                name="price_analysis",
                title="单价分析", 
                data_types=[DataType.UNIT_PRICE, DataType.COST],
            ),
            "comprehensive": PageConfig(
                name="comprehensive",
                title="综合分析",
                data_types=[DataType.COMPREHENSIVE, DataType.QUANTITY_TREND, DataType.UNIT_PRICE],
            )
        }
        
        self.current_page = None
        self.unified_loader = unified_loader
    
    def render_page_selector(self):
        """渲染页面选择器"""
        page_options = {}
        for page_id, config in self.pages.items():
            page_options[f" {config.title}"] = page_id
        
        selected_display = st.sidebar.selectbox(
            "选择分析页面",
            options=list(page_options.keys()),
            index=0
        )
        
        self.current_page = page_options[selected_display]
        return self.current_page
    
    def get_page_config(self, page_id: str = None) -> PageConfig:
        """获取页面配置"""
        page_id = page_id or self.current_page
        return self.pages.get(page_id)
    
    def load_page_data(self, page_id: str, start_date: date, end_date: date, 
                      filters: Optional[Dict] = None) -> Dict[str, pd.DataFrame]:
        """为指定页面加载所需的所有数据"""
        config = self.get_page_config(page_id)
        if not config:
            return {}
        
        data = {}
        for data_type in config.data_types:
            try:
                additional_filters = []
                inv_number = None
                
                if filters:
                    if filters.get('dept_filter') and data_type in [DataType.COST, DataType.COMPREHENSIVE]:
                        additional_filters.append(f"部門名稱 = '{filters['dept_filter']}'")
                    if filters.get('inv_filter') and data_type in [DataType.COST, DataType.UNIT_PRICE, DataType.COMPREHENSIVE]:
                        if data_type == DataType.UNIT_PRICE:
                            inv_number = filters['inv_filter']
                        else:
                            additional_filters.append(f"品名 = '{filters['inv_filter']}'")
                
                df = self.unified_loader.load_data(
                    data_type, 
                    start_date, 
                    end_date, 
                    additional_filters=additional_filters if additional_filters else None,
                    inv_number=inv_number
                )
                data[data_type.value] = df
                
            except Exception as e:
                st.error(f"加载 {data_type.value} 数据时出错: {str(e)}")
                data[data_type.value] = pd.DataFrame()
        
        return data

# 全局页面管理器实例
page_manager = PageManager()

@st.cache_data
def get_filter_options(df):
    """从DataFrame中获取筛选选项"""
    if df.empty:
        return pd.DataFrame(), pd.DataFrame()
    
    dept_options = df[['DEPT_CODE_STD', 'DEPT_NAME_STD']].drop_duplicates().sort_values('DEPT_NAME_STD')
    inv_options = df[['INV', 'INV_PART_DESCRIPTION']].drop_duplicates().sort_values('INV_PART_DESCRIPTION')
    
    return dept_options, inv_options# =============================================================================
# 分析功能函数
# =============================================================================

def calculate_monthly_summary(df):
    """按月计算汇总数据 - 基于分开查询的逻辑"""
    if df.empty:
        return pd.DataFrame(), pd.DataFrame()
    
    df_work = df.copy()
    
    # 物料级别的汇总数据（已经包含单耗）
    monthly_material_summary = df_work.copy()
    
    # 部门级别的汇总数据
    monthly_dept_summary = df_work.groupby(['YEAR_MONTH', 'DEPT_CODE_STD', 'DEPT_NAME_STD']).agg({
        'TOTAL_COST': 'sum',     # 部门总成本
        'WPNL_SIZE': 'first',    # 部门产出（所有行都相同）
        'QUAN': 'sum',           # 部门总数量
        'INV': 'nunique'         # 部门物料种类数
    }).reset_index()
    
    # 计算部门单耗
    monthly_dept_summary['UNIT_CONSUMPTION'] = monthly_dept_summary['TOTAL_COST'] / monthly_dept_summary['WPNL_SIZE']
    monthly_dept_summary['UNIT_CONSUMPTION'] = monthly_dept_summary['UNIT_CONSUMPTION'].replace([np.inf, -np.inf], 0).fillna(0)
    
    return monthly_material_summary, monthly_dept_summary

def calculate_comprehensive_variance_analysis(df):
    """根据新的业务逻辑进行全面的变差分析"""
    if df.empty:
        return {}
    
    results = {}
    
    # 1. 单耗变差分析（与上月对比）
    if 'DATE' in df.columns:
        # 按物料、部门、日期计算日单耗
        daily_data = df.groupby(['INV', 'DEPT_NAME_STD', 'DATE']).agg({
            'TOTAL_COST': 'sum',
            'QUAN': 'sum'
        }).reset_index()
        
        # 获取部门日产出
        dept_daily_output = df.groupby(['DEPT_NAME_STD', 'DATE'])['WPNL_SIZE'].first().reset_index()
        
        # 合并计算日单耗
        daily_merged = daily_data.merge(dept_daily_output, on=['DEPT_NAME_STD', 'DATE'], how='left')
        daily_merged['日单耗'] = (daily_merged['TOTAL_COST'] / daily_merged['WPNL_SIZE']).replace([np.inf, -np.inf], 0).fillna(0)
        
        # 按月汇总计算月单耗
        daily_merged['年月'] = pd.to_datetime(daily_merged['DATE']).dt.to_period('M')
        monthly_unit_data = daily_merged.groupby(['INV', 'DEPT_NAME_STD', '年月']).agg({
            'TOTAL_COST': 'sum',
            'WPNL_SIZE': 'sum'
        }).reset_index()
        monthly_unit_data['月单耗'] = (monthly_unit_data['TOTAL_COST'] / monthly_unit_data['WPNL_SIZE']).replace([np.inf, -np.inf], 0).fillna(0)
        
        # 计算单耗变差（与上月对比）
        monthly_unit_sorted = monthly_unit_data.sort_values(['INV', 'DEPT_NAME_STD', '年月'])
        monthly_unit_sorted['上月单耗'] = monthly_unit_sorted.groupby(['INV', 'DEPT_NAME_STD'])['月单耗'].shift(1)
        monthly_unit_sorted['单耗变差'] = monthly_unit_sorted['月单耗'] - monthly_unit_sorted['上月单耗']
        monthly_unit_sorted['单耗变差百分比'] = ((monthly_unit_sorted['单耗变差'] / monthly_unit_sorted['上月单耗']) * 100).replace([np.inf, -np.inf], np.nan).round(2)
        
        # 筛选最新月份且单耗增加的记录
        latest_month = monthly_unit_sorted['年月'].max()
        unit_variance = monthly_unit_sorted[
            (monthly_unit_sorted['年月'] == latest_month) & 
            (monthly_unit_sorted['单耗变差百分比'].notna()) &
            (monthly_unit_sorted['单耗变差'] > 0)
        ].copy()
        
        # 排序：先按变差绝对值，再按变差百分比绝对值
        unit_variance['变差绝对值'] = unit_variance['单耗变差'].abs()
        unit_variance['百分比绝对值'] = unit_variance['单耗变差百分比'].abs()
        unit_variance = unit_variance.sort_values(['变差绝对值', '百分比绝对值'], ascending=[False, False])
        
        results['单耗变差'] = unit_variance.head(10)
        results['日单耗数据'] = daily_merged
    
    # 2. 耗用量变差分析（与前一天对比）
    if 'DATE' in df.columns:
        daily_qty = df.groupby(['INV', 'DATE'])['QUAN'].sum().reset_index()
        daily_qty_sorted = daily_qty.sort_values(['INV', 'DATE'])
        daily_qty_sorted['前一天数量'] = daily_qty_sorted.groupby('INV')['QUAN'].shift(1)
        daily_qty_sorted['数量变差'] = daily_qty_sorted['QUAN'] - daily_qty_sorted['前一天数量']
        daily_qty_sorted['数量变差百分比'] = ((daily_qty_sorted['数量变差'] / daily_qty_sorted['前一天数量']) * 100).replace([np.inf, -np.inf], np.nan).round(2)
        
        # 筛选最新日期且数量增加的记录
        latest_date = daily_qty_sorted['DATE'].max()
        qty_variance = daily_qty_sorted[
            (daily_qty_sorted['DATE'] == latest_date) & 
            (daily_qty_sorted['数量变差百分比'].notna()) &
            (daily_qty_sorted['数量变差'] > 0)
        ].copy()
        
        # 排序：先按变差绝对值，再按变差百分比绝对值
        qty_variance['变差绝对值'] = qty_variance['数量变差'].abs()
        qty_variance['百分比绝对值'] = qty_variance['数量变差百分比'].abs()
        qty_variance = qty_variance.sort_values(['变差绝对值', '百分比绝对值'], ascending=[False, False])
        
        results['数量变差'] = qty_variance.head(10)
        results['每日数量数据'] = daily_qty_sorted
    
    # 3. 月耗用量变差分析（与上月对比）
    df_copy = df.copy()
    df_copy['年月'] = pd.to_datetime(df_copy['YYYYMM'] + '01', format='%Y%m%d').dt.to_period('M')
    monthly_qty = df_copy.groupby(['INV', '年月'])['QUAN'].sum().reset_index()
    monthly_qty_sorted = monthly_qty.sort_values(['INV', '年月'])
    monthly_qty_sorted['上月数量'] = monthly_qty_sorted.groupby('INV')['QUAN'].shift(1)
    monthly_qty_sorted['月数量变差'] = monthly_qty_sorted['QUAN'] - monthly_qty_sorted['上月数量']
    monthly_qty_sorted['月数量变差百分比'] = ((monthly_qty_sorted['月数量变差'] / monthly_qty_sorted['上月数量']) * 100).replace([np.inf, -np.inf], np.nan).round(2)
    
    # 筛选最新月份且数量增加的记录
    latest_month = monthly_qty_sorted['年月'].max()
    monthly_qty_variance = monthly_qty_sorted[
        (monthly_qty_sorted['年月'] == latest_month) & 
        (monthly_qty_sorted['月数量变差百分比'].notna()) &
        (monthly_qty_sorted['月数量变差'] > 0)
    ].copy()
    
    # 排序：先按变差绝对值，再按变差百分比绝对值
    monthly_qty_variance['变差绝对值'] = monthly_qty_variance['月数量变差'].abs()
    monthly_qty_variance['百分比绝对值'] = monthly_qty_variance['月数量变差百分比'].abs()
    monthly_qty_variance = monthly_qty_variance.sort_values(['变差绝对值', '百分比绝对值'], ascending=[False, False])
    
    results['月数量变差'] = monthly_qty_variance.head(10)
    
    return results

def render_analysis_results(df):
    """渲染分析结果 - 更新版本"""
    if df.empty:
        st.warning("没有数据可供分析")
        return
    
    # 计算汇总数据和变差分析
    monthly_material_summary, monthly_dept_summary = calculate_monthly_summary(df)
    variance_analysis = calculate_comprehensive_variance_analysis(df)
    
    # 显示关键指标
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        total_cost = df['TOTAL_COST'].sum()
        st.metric("总成本", f"¥{total_cost/10000:.2f}万")
    
    with col2:
        total_quan = df['QUAN'].sum()
        st.metric("总数量", f"{total_quan:,.0f}")
    
    with col3:
        # 由于产出是部门级别的，需要去重计算
        total_wpnl = df.groupby(['DEPT_NAME_STD', 'YYYYMM'])['WPNL_SIZE'].first().sum()
        st.metric("总产出", f"{total_wpnl/10000:.2f}万")
    
    with col4:
        # 计算总体平均单耗
        total_unit_consumption = total_cost / total_wpnl if total_wpnl > 0 else 0
        st.metric("总体平均单耗", f"¥{total_unit_consumption:.2f}")
    
    # 创建标签页 - 根据新的分析框架重新设计
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["数量分析", "部门分析", "物料分析", "单价分析", "月度趋势","AI分析"])
    
    with tab1:
        st.subheader("数量分析 - 每个物料的每日领用量")
        
        # 显示变差最大物料的数量趋势图
        if '数量变差' in variance_analysis and not variance_analysis['数量变差'].empty:
            # 获取变差最大的前10个物料
            top_variance_materials = variance_analysis['数量变差']['INV'].tolist()
            
            if '每日数量数据' in variance_analysis:
                # 筛选变差最大的物料数据用于绘图
                plot_data = variance_analysis['每日数量数据'][
                    variance_analysis['每日数量数据']['INV'].isin(top_variance_materials)
                ]
                
                if not plot_data.empty:
                    # 绘制变差最大物料的每日数量趋势
                    fig_qty_trend = px.line(
                        plot_data,
                        x='DATE',
                        y='QUAN',
                        color='INV',
                        title='数量变差最大的前10个物料每日领用量趋势',
                        labels={'QUAN': '领用量', 'DATE': '日期', 'INV': '物料'}
                    )
                    fig_qty_trend.update_layout(height=500)
                    st.plotly_chart(fig_qty_trend, use_container_width=True)
        
        # 每日耗用量变差前10名明细
        st.subheader("每日耗用量变差前10名明细（较前一天）")
        
        if '数量变差' in variance_analysis and not variance_analysis['数量变差'].empty:
            qty_variance_display = variance_analysis['数量变差'][
                ['INV', 'DATE', 'QUAN', '前一天数量', '数量变差', '数量变差百分比']
            ].copy()
            qty_variance_display.columns = ['物料', '当前日期', '当日数量', '前一天数量', '数量变差', '变差百分比(%)']
            
            # 处理边界情况：当前一天数量为0时显示特殊标记
            qty_variance_display['变差百分比(%)'] = qty_variance_display['变差百分比(%)'].fillna('∞')
            
            try:
                st.dataframe(
                    qty_variance_display.style.format({
                        '当日数量': '{:,.0f}',
                        '前一天数量': '{:,.0f}',
                        '数量变差': '{:+,.0f}',
                        '变差百分比(%)': lambda x: f'{x:+.2f}%' if isinstance(x, (int, float)) else str(x)
                    }),
                    use_container_width=True
                )
            except:
                st.dataframe(qty_variance_display, use_container_width=True)
        else:
            st.info("没有足够的数据进行每日耗用量变差分析")
        
        # 月度耗用量变差前10名明细
        st.subheader("月度耗用量变差前10名明细（较上月）")
        
        if '月数量变差' in variance_analysis and not variance_analysis['月数量变差'].empty:
            monthly_qty_variance_display = variance_analysis['月数量变差'][
                ['INV', '年月', 'QUAN', '上月数量', '月数量变差', '月数量变差百分比']
            ].copy()
            monthly_qty_variance_display.columns = ['物料', '当前月份', '当月数量', '上月数量', '数量变差', '变差百分比(%)']
            
            # 处理边界情况
            monthly_qty_variance_display['变差百分比(%)'] = monthly_qty_variance_display['变差百分比(%)'].fillna('∞')
            
            try:
                st.dataframe(
                    monthly_qty_variance_display.style.format({
                        '当月数量': '{:,.0f}',
                        '上月数量': '{:,.0f}',
                        '数量变差': '{:+,.0f}',
                        '变差百分比(%)': lambda x: f'{x:+.2f}%' if isinstance(x, (int, float)) else str(x)
                    }),
                    use_container_width=True
                )
            except:
                st.dataframe(monthly_qty_variance_display, use_container_width=True)
        else:
            st.info("没有足够的数据进行月度耗用量变差分析")
    
    with tab2:
        st.subheader("部门分析 - 单耗排名")
        
        # 部门单耗排名
        if '日单耗数据' in variance_analysis:
            # 计算各部门内物料的平均单耗排名
            dept_material_unit_consumption = variance_analysis['日单耗数据'].groupby(['DEPT_NAME_STD', 'INV']).agg({
                '日单耗': 'mean',
                'TOTAL_COST': 'sum'
            }).reset_index()
            
            # 选择部门进行分析
            dept_list = dept_material_unit_consumption['DEPT_NAME_STD'].unique()
            selected_dept = st.selectbox("选择部门查看单耗排名：", dept_list)
            
            if selected_dept:
                dept_data = dept_material_unit_consumption[
                    dept_material_unit_consumption['DEPT_NAME_STD'] == selected_dept
                ].sort_values('日单耗', ascending=False)
                
                st.subheader(f"{selected_dept} - 物料单耗排名")
                
                # 显示排名表
                dept_ranking_display = dept_data[['INV', '日单耗', 'TOTAL_COST']].copy()
                dept_ranking_display['总成本'] = dept_ranking_display['TOTAL_COST'] / 10000  # 转换为万元
                dept_ranking_display.columns = ['物料', '平均日单耗', '总成本']
                dept_ranking_display['排名'] = range(1, len(dept_ranking_display) + 1)
                dept_ranking_display = dept_ranking_display[['排名', '物料', '平均日单耗', '总成本']]
                
                try:
                    st.dataframe(
                        dept_ranking_display.style.format({
                            '平均日单耗': '¥{:.2f}',
                            '总成本': '¥{:.2f}万'
                        }),
                        use_container_width=True
                    )
                except:
                    st.dataframe(dept_ranking_display, use_container_width=True)
                
                # 显示该部门的单耗变差
                if '单耗变差' in variance_analysis:
                    dept_unit_variance = variance_analysis['单耗变差'][
                        variance_analysis['单耗变差']['DEPT_NAME_STD'] == selected_dept
                    ]
                    
                    if not dept_unit_variance.empty:
                        st.subheader(f"{selected_dept} - 单耗变差前10名明细（较上月）")
                        
                        dept_variance_display = dept_unit_variance[
                            ['INV', '年月', '月单耗', '上月单耗', '单耗变差', '单耗变差百分比']
                        ].copy()
                        dept_variance_display.columns = ['物料', '当前月份', '本月单耗', '上月单耗', '单耗变差', '变差百分比(%)']
                        
                        # 处理边界情况
                        dept_variance_display['变差百分比(%)'] = dept_variance_display['变差百分比(%)'].fillna('∞')
                        
                        try:
                            st.dataframe(
                                dept_variance_display.style.format({
                                    '本月单耗': '¥{:.2f}',
                                    '上月单耗': '¥{:.2f}',
                                    '单耗变差': '¥{:+.2f}',
                                    '变差百分比(%)': lambda x: f'{x:+.2f}%' if isinstance(x, (int, float)) else str(x)
                                }),
                                use_container_width=True
                            )
                        except:
                            st.dataframe(dept_variance_display, use_container_width=True)
                    else:
                        st.info(f"{selected_dept} 暂无单耗变差数据")
        else:
            st.info("没有足够的数据进行部门单耗排名分析")
    
    with tab3:
        st.subheader("部门成本分析")
        
        if not monthly_dept_summary.empty:
            # 部门成本排名
            dept_ranking = monthly_dept_summary.groupby(['DEPT_CODE_STD', 'DEPT_NAME_STD']).agg({
                'TOTAL_COST': 'sum',
                'WPNL_SIZE': 'sum',
                'INV': 'sum'  # 物料种类总数
            }).reset_index().sort_values('TOTAL_COST', ascending=False)
            
            # 计算部门单耗
            dept_ranking['UNIT_CONSUMPTION'] = dept_ranking['TOTAL_COST'] / dept_ranking['WPNL_SIZE']
            dept_ranking['UNIT_CONSUMPTION'] = dept_ranking['UNIT_CONSUMPTION'].fillna(0)
            
            # 添加单耗排名
            dept_unit_ranking = dept_ranking.sort_values('UNIT_CONSUMPTION', ascending=False)
            
            col1, col2 = st.columns(2)
            
            with col1:
                # 部门成本排名图
                fig_cost = px.bar(dept_ranking.head(10), 
                       x='DEPT_NAME_STD', 
                       y='TOTAL_COST',
                       title="部门成本排名 (前10名)")
                fig_cost.update_xaxes(tickangle=45)
                st.plotly_chart(fig_cost, use_container_width=True)
            
            with col2:
                # 部门单耗排名图
                fig_unit = px.bar(dept_unit_ranking.head(10), 
                           x='DEPT_NAME_STD', 
                           y='UNIT_CONSUMPTION',
                           title="部门单耗排名 (前10名)",
                           color='UNIT_CONSUMPTION',
                           color_continuous_scale='Reds')
                fig_unit.update_xaxes(tickangle=45)
                st.plotly_chart(fig_unit, use_container_width=True)
            
            # 部门详细表格
            st.subheader("部门汇总数据")
            
            # 检查必要的列是否存在
            required_cols = ['DEPT_NAME_STD', 'TOTAL_COST', 'WPNL_SIZE', 'INV', 'UNIT_CONSUMPTION']
            if all(col in dept_ranking.columns for col in required_cols):
                display_dept_df = dept_ranking[required_cols].copy()
                display_dept_df['TOTAL_COST'] = display_dept_df['TOTAL_COST'] / 10000  # 转换为万元
                display_dept_df['WPNL_SIZE'] = display_dept_df['WPNL_SIZE'] / 10000  # 转换为万
                display_dept_df.columns = ['部门名称', '总成本', '产出', '物料种类数', '部门单耗']
            else:
                st.error(f"数据结构异常，缺少必要的列。现有列：{list(dept_ranking.columns)}")
                display_dept_df = pd.DataFrame()
            
            try:
                st.dataframe(
                    display_dept_df.style.format({
                        '总成本': '¥{:.2f}万',
                        '产出': '{:.2f}万',
                        '物料种类数': '{:,.0f}',
                        '部门单耗': '¥{:.2f}'
                    }),
                    use_container_width=True
                )
            except Exception:
                st.dataframe(display_dept_df, use_container_width=True)
                
            # 单耗排名表
            st.subheader("部门单耗排名")
            required_ranking_cols = ['DEPT_NAME_STD', 'UNIT_CONSUMPTION', 'TOTAL_COST', 'WPNL_SIZE']
            if all(col in dept_unit_ranking.columns for col in required_ranking_cols):
                display_unit_ranking_df = dept_unit_ranking[required_ranking_cols].copy()
                display_unit_ranking_df['TOTAL_COST'] = display_unit_ranking_df['TOTAL_COST'] / 10000  # 转换为万元
                display_unit_ranking_df['WPNL_SIZE'] = display_unit_ranking_df['WPNL_SIZE'] / 10000  # 转换为万
                display_unit_ranking_df.columns = ['部门名称', '单耗', '总成本', '产出']
                display_unit_ranking_df['排名'] = range(1, len(display_unit_ranking_df) + 1)
                display_unit_ranking_df = display_unit_ranking_df[['排名', '部门名称', '单耗', '总成本', '产出']]
            else:
                st.error(f"单耗排名数据结构异常，缺少必要的列。现有列：{list(dept_unit_ranking.columns)}")
                display_unit_ranking_df = pd.DataFrame()
            
            try:
                st.dataframe(
                    display_unit_ranking_df.style.format({
                        '单耗': '¥{:.2f}',
                        '总成本': '¥{:.2f}万',
                        '产出': '{:.2f}万'
                    }),
                    use_container_width=True
                )
            except Exception:
                st.dataframe(display_unit_ranking_df, use_container_width=True)# =============================================================================
# 主要功能函数和界面渲染
# =============================================================================

def render_analysis_results_part2(df, monthly_material_summary, variance_analysis):
    """渲染分析结果的第二部分 - 物料分析、单价分析等"""
    
    # 继续tab4 - 单价分析
    tab4, tab5, tab6 = st.tabs(["物料分析", "月度趋势", "AI分析"])
    
    with tab4:
        st.subheader("单价分析")
        
        if not monthly_material_summary.empty:
            # 按部门分组进行物料分析
            departments = monthly_material_summary['DEPT_NAME_STD'].unique()
            
            # 部门选择器
            selected_dept_analysis = st.selectbox(
                "选择部门进行物料单耗分析",
                options=['全部部门'] + list(departments),
                key='dept_material_analysis'
            )

            if selected_dept_analysis == '全部部门':
                # 显示所有部门的物料成本排名
                inv_ranking = (
                    monthly_material_summary.groupby(
                        ['INV', 'INV_PART_DESCRIPTION', 'DEPT_CODE_STD', 'DEPT_NAME_STD']
                    )
                    .agg({
                        'TOTAL_COST': 'sum',
                        'QUAN': 'sum',
                        'WPNL_SIZE': 'first',  # 产出
                        'UNIT_CONSUMPTION': 'mean'  # 平均单耗
                    })
                    .reset_index()
                    .sort_values('TOTAL_COST', ascending=False)
                )

                fig = px.bar(
                    inv_ranking.head(15),
                    x='INV_PART_DESCRIPTION',
                    y='TOTAL_COST',
                    color='DEPT_NAME_STD',
                    title="物料成本排名 (前15名)"
                )
                fig.update_xaxes(tickangle=45)
                st.plotly_chart(fig, use_container_width=True)

                # 物料详细表格
                st.subheader("物料汇总数据")
                required_inv_cols = [
                    'INV_PART_DESCRIPTION', 'DEPT_NAME_STD', 'TOTAL_COST', 'QUAN', 'WPNL_SIZE', 'UNIT_CONSUMPTION'
                ]
                if all(col in inv_ranking.columns for col in required_inv_cols):
                    display_inv_df = inv_ranking[required_inv_cols].copy()
                    display_inv_df['TOTAL_COST'] = display_inv_df['TOTAL_COST'] / 10000  # 转换为万元
                    display_inv_df['WPNL_SIZE'] = display_inv_df['WPNL_SIZE'] / 10000  # 转换为万
                    display_inv_df.columns = ['物料描述', '所属部门', '总成本', '总数量', '产出', '单耗']
                else:
                    st.error(
                        f"物料数据结构异常，缺少必要的列。现有列：{list(inv_ranking.columns)}"
                    )
                    display_inv_df = pd.DataFrame()

                try:
                    st.dataframe(
                        display_inv_df.style.format({
                            '总成本': '¥{:.2f}万',
                            '总数量': '{:,.0f}',
                            '产出': '{:.2f}万',
                            '单耗': '¥{:.2f}'
                        }),
                        use_container_width=True
                    )
                except Exception:
                    st.dataframe(display_inv_df, use_container_width=True)
            else:
                # 分析特定部门的物料单耗趋势
                dept_materials = monthly_material_summary[
                    monthly_material_summary['DEPT_NAME_STD'] == selected_dept_analysis
                ].copy()

                if not dept_materials.empty:
                    st.subheader(f"{selected_dept_analysis} - 物料单耗趋势分析")
                    
                    # 按物料和年月分析单耗趋势
                    material_trend = dept_materials.groupby(['INV', 'YEAR_MONTH']).agg({
                        'UNIT_CONSUMPTION': 'mean',
                        'TOTAL_COST': 'sum',
                        'QUAN': 'sum'
                    }).reset_index()
            
                    # 获取单耗最高的前10个物料
                    top_materials_unit = dept_materials.groupby('INV')['UNIT_CONSUMPTION'].mean().nlargest(10)
                    top_material_trend = material_trend[material_trend['INV'].isin(top_materials_unit.index)]
                    
                    if not top_material_trend.empty:
                        # 转换年月为字符串
                        top_material_trend['年月字符串'] = top_material_trend['YEAR_MONTH'].astype(str)
                        
                        # 绘制单耗趋势图
                        fig_unit_trend = px.line(
                            top_material_trend,
                            x='年月字符串',
                            y='UNIT_CONSUMPTION',
                            color='INV',
                            title=f'{selected_dept_analysis} - 前10大单耗物料趋势',
                            labels={'UNIT_CONSUMPTION': '单耗 (¥)', '年月字符串': '年月', 'INV': '物料'}
                        )
                        fig_unit_trend.update_layout(height=500)
                        st.plotly_chart(fig_unit_trend, use_container_width=True)
                        
                        # 单耗变差分析
                        st.subheader("物料单耗变差前10名明细")
                        
                        # 计算月度变化
                        material_trend_sorted = material_trend.sort_values(['INV', 'YEAR_MONTH'])
                        material_trend_sorted['上月单耗'] = material_trend_sorted.groupby('INV')['UNIT_CONSUMPTION'].shift(1)
                        material_trend_sorted['单耗变差'] = material_trend_sorted['UNIT_CONSUMPTION'] - material_trend_sorted['上月单耗']
                        material_trend_sorted['单耗变差百分比'] = (material_trend_sorted['单耗变差'] / material_trend_sorted['上月单耗'] * 100).round(2)
                        
                        # 筛选最近月份的数据
                        latest_month = material_trend_sorted['YEAR_MONTH'].max()
                        current_month_unit_changes = material_trend_sorted[
                            (material_trend_sorted['YEAR_MONTH'] == latest_month) & 
                            (material_trend_sorted['单耗变差百分比'].notna())
                        ].copy()
                        
                        current_month_unit_changes['单耗变差绝对值'] = current_month_unit_changes['单耗变差百分比'].abs()
                        top_unit_changes = current_month_unit_changes.nlargest(10, '单耗变差绝对值')
                        
                        if not top_unit_changes.empty:
                            display_unit_changes_df = top_unit_changes[['INV', 'YEAR_MONTH', 'UNIT_CONSUMPTION', '上月单耗', '单耗变差', '单耗变差百分比']].copy()
                            display_unit_changes_df.columns = ['物料', '当前月份', '当月单耗', '上月单耗', '单耗变差', '单耗变差百分比(%)']
                            
                            try:
                                st.dataframe(
                                    display_unit_changes_df.style.format({
                                        '当月单耗': '¥{:.2f}',
                                        '上月单耗': '¥{:.2f}',
                                        '单耗变差': '¥{:+.2f}',
                                        '单耗变差百分比(%)': '{:+.2f}%'
                                    }),
                                    use_container_width=True
                                )
                            except Exception:
                                st.dataframe(display_unit_changes_df, use_container_width=True)
                        else:
                            st.info("没有足够的月度数据进行单耗变差分析")
                    else:
                        st.info("该部门没有物料单耗趋势数据")
                else:
                    st.info("该部门没有物料数据")
        
        # 添加物料单耗趋势对比功能
        st.divider()
        st.subheader("物料单耗趋势对比")
        
        if not monthly_material_summary.empty:
            # 获取所有物料列表
            all_materials = monthly_material_summary['INV'].unique()
            
            # 多选物料控件
            selected_materials = st.multiselect(
                "选择要对比的物料（可选择多个）:",
                options=all_materials,
                default=[],
                help="选择你要对比单耗趋势的物料，最多建议选择8个以便于查看"
            )
            
            if selected_materials:
                # 筛选选中物料的数据
                material_comparison_data = monthly_material_summary[
                    monthly_material_summary['INV'].isin(selected_materials)
                ].copy()
                
                if not material_comparison_data.empty:
                    # 准备趋势图数据
                    material_trend_comparison = material_comparison_data.groupby(['INV', 'YEAR_MONTH']).agg({
                        'UNIT_CONSUMPTION': 'mean',
                        'TOTAL_COST': 'sum',
                        'QUAN': 'sum',
                        'DEPT_NAME_STD': 'first'
                    }).reset_index()
                    
                    # 转换年月为字符串便于显示
                    material_trend_comparison['年月字符串'] = material_trend_comparison['YEAR_MONTH'].astype(str)
                    
                    # 绘制单耗趋势对比图
                    fig_comparison = px.line(
                        material_trend_comparison,
                        x='年月字符串',
                        y='UNIT_CONSUMPTION',
                        color='INV',
                        title='选定物料单耗趋势对比',
                        labels={'UNIT_CONSUMPTION': '单耗 (¥)', '年月字符串': '年月', 'INV': '物料'},
                        markers=True
                    )
                    fig_comparison.update_layout(
                        height=600,
                        hovermode='x unified',
                        legend=dict(
                            orientation="h",
                            yanchor="bottom",
                            y=1.02,
                            xanchor="right",
                            x=1
                        )
                    )
                    st.plotly_chart(fig_comparison, use_container_width=True)
                    
                    # 显示对比数据表
                    st.subheader("物料单耗对比数据详情")
                    comparison_table = material_trend_comparison.pivot_table(
                        index='年月字符串',
                        columns='INV',
                        values='UNIT_CONSUMPTION',
                        aggfunc='mean'
                    ).fillna(0)
                    
                    # 重命名索引
                    comparison_table.index.name = '年月'
                    
                    # 添加平均值行
                    avg_row = comparison_table.mean()
                    avg_row.name = '平均值'
                    comparison_table = pd.concat([comparison_table, avg_row.to_frame().T])
                    
                    try:
                        st.dataframe(
                            comparison_table.style.format('¥{:.2f}'),
                            use_container_width=True
                        )
                    except Exception:
                        st.dataframe(comparison_table, use_container_width=True)
                    
                    # 添加统计摘要
                    st.subheader("统计摘要")
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        # 最高/最低单耗物料
                        avg_consumption = material_trend_comparison.groupby('INV')['UNIT_CONSUMPTION'].mean().sort_values(ascending=False)
                        
                        st.write("**平均单耗排名:**")
                        for i, (material, consumption) in enumerate(avg_consumption.items(), 1):
                            st.write(f"{i}. {material}: ¥{consumption:.2f}")
                    
                    with col2:
                        # 趋势变化分析
                        st.write("**趋势变化分析:**")
                        for material in selected_materials:
                            material_data = material_trend_comparison[material_trend_comparison['INV'] == material].sort_values('YEAR_MONTH')
                            if len(material_data) >= 2:
                                first_value = material_data.iloc[0]['UNIT_CONSUMPTION']
                                last_value = material_data.iloc[-1]['UNIT_CONSUMPTION']
                                change_rate = ((last_value - first_value) / first_value * 100) if first_value != 0 else 0
                                trend_indicator = "📈" if change_rate > 0 else "📉" if change_rate < 0 else "➡️"
                                st.write(f"{trend_indicator} {material}: {change_rate:+.2f}%")
                else:
                    st.warning("选定物料没有数据可供对比")
            else:
                st.info("请选择要对比的物料")
        else:
            st.warning("没有物料数据可供分析")
    
    with tab5:
        st.subheader("单价变差分析")

        # 获取日期范围用于单价趋势查询
        if 'DATE' in df.columns:
            start_date = df['DATE'].min()
            end_date = df['DATE'].max()
        else:
            start_date = pd.to_datetime(df['YYYYMM'].min() + '01', format='%Y%m%d')
            end_date = pd.to_datetime(df['YYYYMM'].max() + '01', format='%Y%m%d') + pd.offsets.MonthEnd(0)

        # 加载所有料号的单价数据
        unit_price_df = load_unit_price_data(start_date, end_date)

        if not unit_price_df.empty:
            # 每日单价变差分析（跟前一天对比）
            st.subheader("每日单价变差前10名明细")
            
            # 计算每日单价变化（跟前一天对比）
            unit_price_daily = unit_price_df.sort_values(['INV', 'DATE'])
            unit_price_daily['前一天单价'] = unit_price_daily.groupby('INV')['UNIT_PRICE'].shift(1)
            unit_price_daily['日单价变差'] = unit_price_daily['UNIT_PRICE'] - unit_price_daily['前一天单价']
            unit_price_daily['日单价变差百分比'] = (unit_price_daily['日单价变差'] / unit_price_daily['前一天单价'] * 100).round(2)
            
            # 筛选最新日期的数据并按变差绝对值排序
            latest_date = unit_price_daily['DATE'].max()
            current_day_price_changes = unit_price_daily[
                (unit_price_daily['DATE'] == latest_date) & 
                (unit_price_daily['日单价变差百分比'].notna())
            ].copy()
            
            current_day_price_changes['单价变差绝对值'] = current_day_price_changes['日单价变差百分比'].abs()
            top_daily_price_changes = current_day_price_changes.nlargest(10, '单价变差绝对值')
            
            if not top_daily_price_changes.empty:
                display_daily_price_changes_df = top_daily_price_changes[['INV', 'DATE', '日单价变差百分比']].copy()
                display_daily_price_changes_df.columns = ['物料', '当前日期', '变差百分比(%)']
                
                try:
                    st.dataframe(
                        display_daily_price_changes_df.style.format({
                            '变差百分比(%)': '{:+.2f}%'
                        }),
                        use_container_width=True
                    )
                except Exception:
                    st.dataframe(display_daily_price_changes_df, use_container_width=True)
            else:
                st.info("没有足够的日度数据进行每日单价变差分析")
        
        # 月度单价变差分析（跟上个月对比）
        st.subheader("月度单价变差前10名明细")
        
        # 添加年月字段
        unit_price_df['年月'] = pd.to_datetime(unit_price_df['DATE']).dt.to_period('M')
        
        # 按物料和月份计算平均单价
        monthly_price = unit_price_df.groupby(['INV', '年月'])['UNIT_PRICE'].mean().reset_index()
        
        # 计算月度单价变化
        monthly_price = monthly_price.sort_values(['INV', '年月'])
        monthly_price['上月单价'] = monthly_price.groupby('INV')['UNIT_PRICE'].shift(1)
        monthly_price['单价变差'] = monthly_price['UNIT_PRICE'] - monthly_price['上月单价']
        monthly_price['单价变差百分比'] = (monthly_price['单价变差'] / monthly_price['上月单价'] * 100).round(2)
        
        # 筛选最近月份的数据
        latest_month = monthly_price['年月'].max()
        current_month_price_changes = monthly_price[
            (monthly_price['年月'] == latest_month) & 
            (monthly_price['单价变差百分比'].notna())
        ].copy()
        
        current_month_price_changes['单价变差绝对值'] = current_month_price_changes['单价变差百分比'].abs()
        top_price_changes = current_month_price_changes.nlargest(10, '单价变差绝对值')
        
        if not top_price_changes.empty:
            display_price_changes_df = top_price_changes[['INV', '年月', '单价变差百分比']].copy()
            display_price_changes_df.columns = ['物料', '当前月份', '变差百分比(%)']
            
            try:
                st.dataframe(
                    display_price_changes_df.style.format({
                        '变差百分比(%)': '{:+.2f}%'
                    }),
                    use_container_width=True
                )
            except Exception:
                st.dataframe(display_price_changes_df, use_container_width=True)
        else:
            st.info("没有足够的月度数据进行单价变差分析")
        
        # 原有的单价波动预警功能
        st.subheader("单价波动预警详情")
        
        # 计算每日单价变化率
        unit_price_df = unit_price_df.sort_values(by=['INV', 'DATE'])
        unit_price_df['PREV_UNIT_PRICE'] = unit_price_df.groupby('INV')['UNIT_PRICE'].shift(1)
        unit_price_df['PRICE_CHANGE_PCT'] = (unit_price_df['UNIT_PRICE'] - unit_price_df['PREV_UNIT_PRICE']) / unit_price_df['PREV_UNIT_PRICE']
        unit_price_df['PRICE_CHANGE_PCT'] = unit_price_df['PRICE_CHANGE_PCT'].replace([np.inf, -np.inf], np.nan).fillna(0)

        # 设置预警阈值 (例如：单价变化超过5%)
        warning_threshold = 0.05
        unit_price_df['IS_WARNING'] = (unit_price_df['PRICE_CHANGE_PCT'].abs() > warning_threshold)

        # 筛选出有预警的料号
        all_invs = unit_price_df['INV'].unique()

        if len(all_invs) > 0:
            st.write("选择料号查看单价趋势：")
            selected_inv = st.selectbox("", all_invs, key='price_inv_select')

            # 检查是否有预警，并显示预警信息
            warning_invs = unit_price_df[unit_price_df['IS_WARNING']]['INV'].unique()
            if selected_inv in warning_invs:
                st.warning(f"料号 {selected_inv} 存在单价波动预警 (变化超过 {warning_threshold:.1%})。")
            else:
                st.info(f"料号 {selected_inv} 当前没有单价波动预警。")

            inv_data = unit_price_df[unit_price_df['INV'] == selected_inv].copy()
            inv_data = inv_data.sort_values('DATE')

            # 绘制单价趋势图
            fig_price = px.line(inv_data, x='DATE', y='UNIT_PRICE', title=f"{selected_inv} 单价趋势")
            
            # 标记预警点
            warning_points = inv_data[inv_data['IS_WARNING']]
            if not warning_points.empty:
                fig_price.add_trace(go.Scatter(
                    mode='markers',
                    x=warning_points['DATE'],
                    y=warning_points['UNIT_PRICE'],
                    marker=dict(color='red', size=10, symbol='star'),
                    name='预警点',
                    hovertext=warning_points.apply(lambda row: f"日期: {row['DATE'].strftime('%Y-%m-%d')}<br>单价: {row['UNIT_PRICE']:.4f}<br>变化: {row['PRICE_CHANGE_PCT']:.2%}", axis=1)
                ))

            fig_price.update_xaxes(title_text="日期")
            fig_price.update_yaxes(title_text="单价")
            st.plotly_chart(fig_price, use_container_width=True)

            # 显示详细预警数据
            st.subheader(f"{selected_inv} 单价波动预警数据")
            display_price_warning_df = inv_data[inv_data['IS_WARNING']][['DATE', 'INV', 'PRICE_CHANGE_PCT']].copy()
            display_price_warning_df.columns = ['日期', '料号', '单价变化率']
            try:
                st.dataframe(
                    display_price_warning_df.style.format({
                        '单价变化率': '{:+.2%}'
                    }),
                    use_container_width=True
                )
            except Exception:
                st.dataframe(display_price_warning_df, use_container_width=True)
        else:
            st.info("当前日期范围内没有料号的单价波动超过预警阈值。")
    else:
        st.warning("无法加载单价数据。")
            
    with tab6:
        st.subheader("AI智能分析助手")
        st.caption("与AI对话，分析当前单耗数据")
        
        # LLM配置
        llm_model_name = "QWQ:32B" 
        llm_base_url = "http://172.18.80.88:8106/v1"
        
        # 记录数限制配置
        col1, col2 = st.columns([1, 3])
        with col1:
            max_rows = st.number_input(
                "分析最大数据量", 
                min_value=10, 
                max_value=1000,
                value=500,
                step=50,
                key="cost_llm_max_rows"
            )
        
        # 创建消息容器
        message_container = st.container()
        # 显示历史消息
        with message_container:
            for i, message in enumerate(st.session_state.cost_llm_messages):
                if message["role"] == "user":
                    # 用户消息显示在右侧
                    col1, col2 = st.columns([1, 3])
                    with col2:
                        st.markdown(f"""
                        <div style="background-color: #e3f2fd; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                            <strong>用户:</strong><br>
                            {message["content"]}
                        </div>
                        """, unsafe_allow_html=True)
                else:
                    # LLM消息显示在左侧
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.markdown(f"""
                        <div style="background-color: #f5f5f5; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                            <strong>AI助手:</strong><br>
                            {message["content"]}
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # 为每条助手回答添加导出按钮
                        if message["content"] and not message["content"].startswith("调用本地 LLM"):
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            filename = f"单耗分析报告_{timestamp}.docx"
                            
                            # 准备分析上下文信息
                            date_range = "未知"
                            if 'DATE' in df.columns:
                                min_date = df['DATE'].min().strftime('%Y-%m-%d') if not df['DATE'].isna().all() else '未知'
                                max_date = df['DATE'].max().strftime('%Y-%m-%d') if not df['DATE'].isna().all() else '未知'
                                date_range = f"{min_date} 至 {max_date}"
                            elif 'YYYYMM' in df.columns:
                                min_ym = df['YYYYMM'].min() if not df['YYYYMM'].isna().all() else '未知'
                                max_ym = df['YYYYMM'].max() if not df['YYYYMM'].isna().all() else '未知'
                                date_range = f"{min_ym} 至 {max_ym}"
                                
                            analysis_context = {
                                'date_range': date_range,
                                'total_records': len(df),
                                'llm_model': llm_model_name,
                                'data_summary': {
                                    '分析记录数': len(df),
                                    '数据列数': len(df.columns),
                                    '总成本': f"¥{df['TOTAL_COST'].sum():,.2f}",
                                    '总数量': f"{df['QUAN'].sum():,.0f}",
                                    '部门数量': df['DEPT_NAME_STD'].nunique(),
                                    '物料种类': df['INV'].nunique()
                                }
                            }
                            
                            try:
                                # 生成Word报告
                                with st.spinner(f"正在生成第{i+1}条回答的报告..."):
                                    user_question = ""
                                    if i > 0 and st.session_state.cost_llm_messages[i-1]["role"] == "user":
                                        user_question = st.session_state.cost_llm_messages[i-1]["content"]
                                    elif i == 0:
                                        user_question = "初始分析"
                                    
                                    doc_io = generate_word_report(user_question, message["content"], analysis_context)
                                
                                # 添加下载按钮
                                st.download_button(
                                    label=f"导出第{i+1}条回答为Word报告",
                                    data=doc_io.getvalue(),
                                    file_name=filename,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"export_word_msg_{i}_{len(st.session_state.cost_llm_messages)}",
                                    help="点击下载Word格式的分析报告",
                                    use_container_width=True
                                )
                                
                            except Exception as export_err:
                                st.error(f"生成报告失败: {export_err}")
        
        # 使用streamlit的chat_input作为输入框
        user_input = st.chat_input("请分析当前单耗数据")
        
        # 处理用户输入
        if user_input:
            # 添加用户消息到会话状态
            st.session_state.cost_llm_messages.append({"role": "user", "content": user_input})
            
            # 显示用户消息
            with message_container:
                col1, col2 = st.columns([1, 3])
                with col2:
                    st.markdown(f"""
                    <div style="background-color: #e3f2fd; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                        <strong>用户:</strong><br>
                        {user_input}
                    </div>
                    """, unsafe_allow_html=True)
            
            # 显示助手消息
            with message_container:
                col1, col2 = st.columns([3, 1])
                with col1:
                    # 创建一个占位符用于流式输出
                    message_placeholder = st.empty()
                    message_placeholder.markdown("""
                    <div style="background-color: #f5f5f5; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                        <strong>AI助手:</strong><br>
                        正在分析中...
                    </div>
                    """, unsafe_allow_html=True)

                    try:
                        # 创建一个专用于LLM的数据副本
                        llm_df = df.copy()
                        
                        # 需要删除的列列表（分析生成的技术字段）
                        columns_to_remove = [
                            'IS_ANOMALY',
                            'Predicted', 
                            'Lower', 
                            'Upper', 
                            'Actual'
                        ]

                        # 从DataFrame中删除存在的列
                        for col in columns_to_remove:
                            if col in llm_df.columns:
                                llm_df.drop(col, axis=1, inplace=True)
                        
                        # 处理DATE列的格式转换和排序
                        if 'DATE' in llm_df.columns:
                            try:
                                # 确保DATE是datetime类型用于排序
                                llm_df['DATE'] = pd.to_datetime(llm_df['DATE'])
                                # 按时间从新到旧排序
                                llm_df = llm_df.sort_values('DATE', ascending=False)
                                # 限制为最新的max_rows条记录
                                llm_df_final = llm_df.head(max_rows)
                                # 将DATE转为字符串格式用于JSON序列化
                                llm_df_final['DATE'] = llm_df_final['DATE'].dt.strftime('%Y-%m-%d')
                            except Exception as date_err:
                                st.warning(f"处理DATE列时出错: {date_err}. 将使用未排序的数据。")
                                llm_df_final = llm_df.head(max_rows)
                        else:
                            # 按YYYYMM排序
                            try:
                                llm_df = llm_df.sort_values('YYYYMM', ascending=False)
                                llm_df_final = llm_df.head(max_rows)
                            except:
                                llm_df_final = llm_df.head(max_rows)
                        
                        # 将精简后的数据集转换为CSV格式
                        data_context = llm_df_final.to_csv(index=False)
                        
                        # 调用LLM
                        try:
                            system_prompt = """你是一个专业的成本分析师和数据分析助手。请基于提供的单耗分析数据和用户问题进行深入分析并回答。

                                数据字段说明：
                                - YYYYMM: 年月
                                - DEPT_NAME_STD: 部门名称
                                - INV: 物料品名
                                - INV_PART_DESCRIPTION: 物料描述
                                - QUAN: 数量
                                - STD_COST: 单价
                                - TOTAL_COST: 总成本
                                - WPNL_SIZE: 产出
                                - UNIT_CONSUMPTION: 单耗（总成本/产出）
                                - YEAR_MONTH: 年月期间
                                - DEPT_CODE_STD: 部门代码

                                请根据数据特点提供专业的成本分析建议。"""
                            
                            data_description = f"单耗分析数据：最新的{len(llm_df_final)}条记录\n数据时间范围：{llm_df_final.get('YYYYMM', pd.Series()).min() or '未知'} 至 {llm_df_final.get('YYYYMM', pd.Series()).max() or '未知'}\n\n"
                            
                            llm = ChatOpenAI(
                                model=llm_model_name,
                                openai_api_key="dummy-key",
                                openai_api_base=llm_base_url,
                                temperature=0,
                                streaming=True  # 启用流式输出
                            )
                            
                            enhanced_prompt = f"""
                                {data_description}以下是CSV格式的单耗分析数据：
                                ```csv
                                {data_context}
                                ```

                                用户问题: {user_input}

                                请根据以上单耗分析数据进行专业分析并回答。
                                """
                            
                            messages = [
                                SystemMessage(content=system_prompt),
                                HumanMessage(content=enhanced_prompt)
                            ]
                            
                            # 使用流式输出
                            llm_response_content = ""
                            for chunk in llm.stream(messages):
                                if chunk.content:
                                    llm_response_content += chunk.content
                                    message_placeholder.markdown(f"""
                                    <div style="background-color: #f5f5f5; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                                        <strong>AI助手:</strong><br>
                                        {llm_response_content}▌
                                    </div>
                                    """, unsafe_allow_html=True)
                            
                            # 移除光标并显示最终内容
                            message_placeholder.markdown(f"""
                            <div style="background-color: #f5f5f5; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                                <strong>AI助手:</strong><br>
                                {llm_response_content}
                            </div>
                            """, unsafe_allow_html=True)
                            
                        except Exception as llm_err:
                            llm_response_content = f"调用本地 LLM (模型: {llm_model_name} @ {llm_base_url}) 时出错: {llm_err}"
                            st.error(llm_response_content)
                            message_placeholder.markdown(f"""
                            <div style="background-color: #f5f5f5; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                                <strong>AI助手:</strong><br>
                                {llm_response_content}
                            </div>
                            """, unsafe_allow_html=True)

                    except Exception as e:
                        error_msg = f"处理数据时出错: {e}"
                        message_placeholder.markdown(f"""
                        <div style="background-color: #f5f5f5; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                            <strong>AI助手:</strong><br>
                            {error_msg}
                        </div>
                        """, unsafe_allow_html=True)
                        llm_response_content = error_msg
                    
                    # 添加助手消息到会话状态
                    st.session_state.cost_llm_messages.append({"role": "assistant", "content": llm_response_content})

def render_sidebar_filters():
    """渲染侧边栏筛选器"""
    with st.sidebar:
        # 数据库连接状态
        with st.status("连接数据库...", expanded=False) as status:
            try:
                with get_db_connection() as conn:
                    cursor = conn.cursor()
                    cursor.execute("SELECT 1 FROM DUAL")
                    cursor.fetchone()
                status.update(label="已连接数据库", state="complete", expanded=False)
            except Exception as e:
                status.update(label=f"数据库连接失败: {e}", state="error", expanded=True)
                st.stop()
        
        st.markdown("---")
        
        # 设置默认时间为最近三个月
        today = date.today()
        default_start = today - timedelta(days=90)  # 三个月前
        default_end = today
        
        # 日期输入控件
        col_start, col_end = st.columns(2)
        with col_start:
            start_date = st.date_input(
                "开始日期", 
                value=default_start
            )
        with col_end:
            end_date = st.date_input(
                "结束日期", 
                value=default_end
            )
         
        # 根据时间范围自动加载筛选选项
        with st.spinner("正在加载筛选选项..."):
            try:
                dept_options, inv_options = load_filter_options(start_date, end_date)
            except Exception as e:
                st.error(f"加载筛选选项失败: {e}")
                dept_options, inv_options = pd.DataFrame(), pd.DataFrame()
        
        # 初始化筛选选项
        selected_dept = "全部"
        selected_inv = "全部"
        
        if not dept_options.empty:
            # 部门筛选
            dept_list = ["全部"] + dept_options['DEPT_CODE_STD'].tolist()
            dept_format_func = lambda x: "全部" if x == "全部" else dept_options[dept_options['DEPT_CODE_STD']==x]['DEPT_NAME_STD'].iloc[0]
            
            selected_dept = st.selectbox(
                "部门选择",
                options=dept_list,
                format_func=dept_format_func,
                help=f"当前时间范围内共有 {len(dept_options)} 个部门"
            )
            
            # 根据选择的部门获取物料选项
            if selected_dept != "全部":
                try:
                    # 加载选定部门下的物料选项
                    with get_db_connection() as conn:
                        query = """
                        SELECT DISTINCT 品名 as INV
                        FROM VI_Wlcb_cost_ai
                        WHERE YYYYMM >= TO_CHAR(:start_date, 'YYYYMM')
                          AND YYYYMM <= TO_CHAR(:end_date, 'YYYYMM')
                          AND 部門名稱 = :dept
                        ORDER BY 品名
                        """
                        df_materials = pd.read_sql(query, conn, params={
                            'start_date': start_date,
                            'end_date': end_date,
                            'dept': selected_dept
                        })
                        if not df_materials.empty:
                            df_materials['INV_PART_DESCRIPTION'] = df_materials['INV']  # 使用品名作为显示名称
                            filtered_inv_options = df_materials
                        else:
                            filtered_inv_options = pd.DataFrame()
                except Exception as e:
                    st.warning(f"加载物料选项失败: {e}")
                    filtered_inv_options = inv_options
            else:
                filtered_inv_options = inv_options
            
            # 物料筛选
            if not filtered_inv_options.empty:
                inv_list = ["全部"] + filtered_inv_options['INV'].tolist()
                inv_format_func = lambda x: "全部" if x == "全部" else filtered_inv_options[filtered_inv_options['INV']==x]['INV_PART_DESCRIPTION'].iloc[0]
                
                material_help_text = f"当前部门下共有 {len(filtered_inv_options)} 个物料" if selected_dept != "全部" else f"所有部门下共有 {len(filtered_inv_options)} 个物料"
                selected_inv = st.selectbox(
                    "物料选择",
                    options=inv_list,
                    format_func=inv_format_func,
                    help=material_help_text
                )
            else:
                st.selectbox("物料选择", ["指定条件下无数据"], disabled=True)
        else:
            # 没有筛选选项时显示提示
            st.selectbox("部门选择", ["指定时间范围内无数据"], disabled=True)
            st.selectbox("物料选择", ["指定时间范围内无数据"], disabled=True)
            st.warning("指定时间范围内没有数据，请调整时间范围")
        
        # 按钮区域
        col_btn1, col_btn2 = st.columns(2)
        
        with col_btn1:
            # 加载数据按钮
            load_data = st.button("加载数据", use_container_width=True, type="primary", disabled=dept_options.empty)
        
        with col_btn2:
            # 重置筛选按钮
            if st.button("重置筛选", use_container_width=True):
                st.session_state.reset_filters = True
                st.rerun()
        
        return {
            'start_date': start_date,
            'end_date': end_date,
            'load_data': load_data,
            'selected_dept': selected_dept,
            'selected_inv': selected_inv,
            'dept_options': dept_options,
            'inv_options': inv_options
        }

def apply_filters_to_data(df, filters):
    """应用筛选条件到数据"""
    if df.empty:
        return df
    
    filtered_df = df.copy()
    
    # 部门筛选 - 使用DEPT_NAME_STD字段
    if filters['selected_dept'] != "全部":
        filtered_df = filtered_df[filtered_df['DEPT_NAME_STD'] == filters['selected_dept']]
    
    # 物料筛选
    if filters['selected_inv'] != "全部":
        filtered_df = filtered_df[filtered_df['INV'] == filters['selected_inv']]
    
    return filtered_df# =============================================================================
# UI渲染和主要功能函数
# =============================================================================

def render_analysis_results(df):
    """渲染分析结果 - 主要入口函数"""
    if df.empty:
        st.warning("没有数据可供分析")
        return
    
    # 计算汇总数据和变差分析
    monthly_material_summary, monthly_dept_summary = calculate_monthly_summary(df)
    variance_analysis = calculate_comprehensive_variance_analysis(df)
    
    # 显示关键指标
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        total_cost = df['TOTAL_COST'].sum()
        st.metric("总成本", f"¥{total_cost/10000:.2f}万")
    
    with col2:
        total_quan = df['QUAN'].sum()
        st.metric("总数量", f"{total_quan:,.0f}")
    
    with col3:
        # 由于产出是部门级别的，需要去重计算
        total_wpnl = df.groupby(['DEPT_NAME_STD', 'YYYYMM'])['WPNL_SIZE'].first().sum()
        st.metric("总产出", f"{total_wpnl/10000:.2f}万")
    
    with col4:
        # 计算总体平均单耗
        total_unit_consumption = total_cost / total_wpnl if total_wpnl > 0 else 0
        st.metric("总体平均单耗", f"¥{total_unit_consumption:.2f}")
    
    # 创建标签页
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["数量分析", "部门分析", "物料分析", "单价分析", "月度趋势","AI分析"])
    
    # 渲染各个标签页
    render_quantity_analysis_tab(tab1, variance_analysis)
    render_department_analysis_tab(tab2, variance_analysis)
    render_material_analysis_tab(tab3, monthly_dept_summary)
    render_price_analysis_tab(tab4, monthly_material_summary, df)
    render_trend_analysis_tab(tab5, df)
    render_ai_analysis_tab(tab6, df)

def render_quantity_analysis_tab(tab, variance_analysis):
    """渲染数量分析标签页"""
    with tab:
        st.subheader("数量分析 - 每个物料的每日领用量")
        
        # 显示变差最大物料的数量趋势图
        if '数量变差' in variance_analysis and not variance_analysis['数量变差'].empty:
            # 获取变差最大的前10个物料
            top_variance_materials = variance_analysis['数量变差']['INV'].tolist()
            
            if '每日数量数据' in variance_analysis:
                # 筛选变差最大的物料数据用于绘图
                plot_data = variance_analysis['每日数量数据'][
                    variance_analysis['每日数量数据']['INV'].isin(top_variance_materials)
                ]
                
                if not plot_data.empty:
                    # 绘制变差最大物料的每日数量趋势
                    fig_qty_trend = px.line(
                        plot_data,
                        x='DATE',
                        y='QUAN',
                        color='INV',
                        title='数量变差最大的前10个物料每日领用量趋势',
                        labels={'QUAN': '领用量', 'DATE': '日期', 'INV': '物料'}
                    )
                    fig_qty_trend.update_layout(height=500)
                    st.plotly_chart(fig_qty_trend, use_container_width=True)
        
        # 每日耗用量变差前10名明细
        st.subheader("每日耗用量变差前10名明细（较前一天）")
        
        if '数量变差' in variance_analysis and not variance_analysis['数量变差'].empty:
            qty_variance_display = variance_analysis['数量变差'][
                ['INV', 'DATE', 'QUAN', '前一天数量', '数量变差', '数量变差百分比']
            ].copy()
            qty_variance_display.columns = ['物料', '当前日期', '当日数量', '前一天数量', '数量变差', '变差百分比(%)']
            
            # 处理边界情况：当前一天数量为0时显示特殊标记
            qty_variance_display['变差百分比(%)'] = qty_variance_display['变差百分比(%)'].fillna('∞')
            
            try:
                st.dataframe(
                    qty_variance_display.style.format({
                        '当日数量': '{:,.0f}',
                        '前一天数量': '{:,.0f}',
                        '数量变差': '{:+,.0f}',
                        '变差百分比(%)': lambda x: f'{x:+.2f}%' if isinstance(x, (int, float)) else str(x)
                    }),
                    use_container_width=True
                )
            except:
                st.dataframe(qty_variance_display, use_container_width=True)
        else:
            st.info("没有足够的数据进行每日耗用量变差分析")
        
        # 月度耗用量变差前10名明细
        st.subheader("月度耗用量变差前10名明细（较上月）")
        
        if '月数量变差' in variance_analysis and not variance_analysis['月数量变差'].empty:
            monthly_qty_variance_display = variance_analysis['月数量变差'][
                ['INV', '年月', 'QUAN', '上月数量', '月数量变差', '月数量变差百分比']
            ].copy()
            monthly_qty_variance_display.columns = ['物料', '当前月份', '当月数量', '上月数量', '数量变差', '变差百分比(%)']
            
            # 处理边界情况
            monthly_qty_variance_display['变差百分比(%)'] = monthly_qty_variance_display['变差百分比(%)'].fillna('∞')
            
            try:
                st.dataframe(
                    monthly_qty_variance_display.style.format({
                        '当月数量': '{:,.0f}',
                        '上月数量': '{:,.0f}',
                        '数量变差': '{:+,.0f}',
                        '变差百分比(%)': lambda x: f'{x:+.2f}%' if isinstance(x, (int, float)) else str(x)
                    }),
                    use_container_width=True
                )
            except:
                st.dataframe(monthly_qty_variance_display, use_container_width=True)
        else:
            st.info("没有足够的数据进行月度耗用量变差分析")

def render_department_analysis_tab(tab, variance_analysis):
    """渲染部门分析标签页"""
    with tab:
        st.subheader("部门分析 - 单耗排名")
        
        # 部门单耗排名
        if '日单耗数据' in variance_analysis:
            # 计算各部门内物料的平均单耗排名
            dept_material_unit_consumption = variance_analysis['日单耗数据'].groupby(['DEPT_NAME_STD', 'INV']).agg({
                '日单耗': 'mean',
                'TOTAL_COST': 'sum'
            }).reset_index()
            
            # 选择部门进行分析
            dept_list = dept_material_unit_consumption['DEPT_NAME_STD'].unique()
            selected_dept = st.selectbox("选择部门查看单耗排名：", dept_list)
            
            if selected_dept:
                dept_data = dept_material_unit_consumption[
                    dept_material_unit_consumption['DEPT_NAME_STD'] == selected_dept
                ].sort_values('日单耗', ascending=False)
                
                st.subheader(f"{selected_dept} - 物料单耗排名")
                
                # 显示排名表
                dept_ranking_display = dept_data[['INV', '日单耗', 'TOTAL_COST']].copy()
                dept_ranking_display['总成本'] = dept_ranking_display['TOTAL_COST'] / 10000  # 转换为万元
                dept_ranking_display.columns = ['物料', '平均日单耗', '总成本']
                dept_ranking_display['排名'] = range(1, len(dept_ranking_display) + 1)
                dept_ranking_display = dept_ranking_display[['排名', '物料', '平均日单耗', '总成本']]
                
                try:
                    st.dataframe(
                        dept_ranking_display.style.format({
                            '平均日单耗': '¥{:.2f}',
                            '总成本': '¥{:.2f}万'
                        }),
                        use_container_width=True
                    )
                except:
                    st.dataframe(dept_ranking_display, use_container_width=True)
                
                # 显示该部门的单耗变差
                if '单耗变差' in variance_analysis:
                    dept_unit_variance = variance_analysis['单耗变差'][
                        variance_analysis['单耗变差']['DEPT_NAME_STD'] == selected_dept
                    ]
                    
                    if not dept_unit_variance.empty:
                        st.subheader(f"{selected_dept} - 单耗变差前10名明细（较上月）")
                        
                        dept_variance_display = dept_unit_variance[
                            ['INV', '年月', '月单耗', '上月单耗', '单耗变差', '单耗变差百分比']
                        ].copy()
                        dept_variance_display.columns = ['物料', '当前月份', '本月单耗', '上月单耗', '单耗变差', '变差百分比(%)']
                        
                        # 处理边界情况
                        dept_variance_display['变差百分比(%)'] = dept_variance_display['变差百分比(%)'].fillna('∞')
                        
                        try:
                            st.dataframe(
                                dept_variance_display.style.format({
                                    '本月单耗': '¥{:.2f}',
                                    '上月单耗': '¥{:.2f}',
                                    '单耗变差': '¥{:+.2f}',
                                    '变差百分比(%)': lambda x: f'{x:+.2f}%' if isinstance(x, (int, float)) else str(x)
                                }),
                                use_container_width=True
                            )
                        except:
                            st.dataframe(dept_variance_display, use_container_width=True)
                    else:
                        st.info(f"{selected_dept} 暂无单耗变差数据")
        else:
            st.info("没有足够的数据进行部门单耗排名分析")

def render_material_analysis_tab(tab, monthly_dept_summary):
    """渲染物料分析标签页"""
    with tab:
        st.subheader("部门成本分析")
        
        if not monthly_dept_summary.empty:
            # 部门成本排名
            dept_ranking = monthly_dept_summary.groupby(['DEPT_CODE_STD', 'DEPT_NAME_STD']).agg({
                'TOTAL_COST': 'sum',
                'WPNL_SIZE': 'sum',
                'INV': 'sum'  # 物料种类总数
            }).reset_index().sort_values('TOTAL_COST', ascending=False)
            
            # 计算部门单耗
            dept_ranking['UNIT_CONSUMPTION'] = dept_ranking['TOTAL_COST'] / dept_ranking['WPNL_SIZE']
            dept_ranking['UNIT_CONSUMPTION'] = dept_ranking['UNIT_CONSUMPTION'].fillna(0)
            
            # 添加单耗排名
            dept_unit_ranking = dept_ranking.sort_values('UNIT_CONSUMPTION', ascending=False)
            
            col1, col2 = st.columns(2)
            
            with col1:
                # 部门成本排名图
                fig_cost = px.bar(dept_ranking.head(10), 
                       x='DEPT_NAME_STD', 
                       y='TOTAL_COST',
                       title="部门成本排名 (前10名)")
                fig_cost.update_xaxes(tickangle=45)
                st.plotly_chart(fig_cost, use_container_width=True)
            
            with col2:
                # 部门单耗排名图
                fig_unit = px.bar(dept_unit_ranking.head(10), 
                           x='DEPT_NAME_STD', 
                           y='UNIT_CONSUMPTION',
                           title="部门单耗排名 (前10名)",
                           color='UNIT_CONSUMPTION',
                           color_continuous_scale='Reds')
                fig_unit.update_xaxes(tickangle=45)
                st.plotly_chart(fig_unit, use_container_width=True)
            
            # 部门详细表格
            st.subheader("部门汇总数据")
            
            # 检查必要的列是否存在
            required_cols = ['DEPT_NAME_STD', 'TOTAL_COST', 'WPNL_SIZE', 'INV', 'UNIT_CONSUMPTION']
            if all(col in dept_ranking.columns for col in required_cols):
                display_dept_df = dept_ranking[required_cols].copy()
                display_dept_df['TOTAL_COST'] = display_dept_df['TOTAL_COST'] / 10000  # 转换为万元
                display_dept_df['WPNL_SIZE'] = display_dept_df['WPNL_SIZE'] / 10000  # 转换为万
                display_dept_df.columns = ['部门名称', '总成本', '产出', '物料种类数', '部门单耗']
            else:
                st.error(f"数据结构异常，缺少必要的列。现有列：{list(dept_ranking.columns)}")
                display_dept_df = pd.DataFrame()
            
            try:
                st.dataframe(
                    display_dept_df.style.format({
                        '总成本': '¥{:.2f}万',
                        '产出': '{:.2f}万',
                        '物料种类数': '{:,.0f}',
                        '部门单耗': '¥{:.2f}'
                    }),
                    use_container_width=True
                )
            except Exception:
                st.dataframe(display_dept_df, use_container_width=True)
                
            # 单耗排名表
            st.subheader("部门单耗排名")
            required_ranking_cols = ['DEPT_NAME_STD', 'UNIT_CONSUMPTION', 'TOTAL_COST', 'WPNL_SIZE']
            if all(col in dept_unit_ranking.columns for col in required_ranking_cols):
                display_unit_ranking_df = dept_unit_ranking[required_ranking_cols].copy()
                display_unit_ranking_df['TOTAL_COST'] = display_unit_ranking_df['TOTAL_COST'] / 10000  # 转换为万元
                display_unit_ranking_df['WPNL_SIZE'] = display_unit_ranking_df['WPNL_SIZE'] / 10000  # 转换为万
                display_unit_ranking_df.columns = ['部门名称', '单耗', '总成本', '产出']
                display_unit_ranking_df['排名'] = range(1, len(display_unit_ranking_df) + 1)
                display_unit_ranking_df = display_unit_ranking_df[['排名', '部门名称', '单耗', '总成本', '产出']]
            else:
                st.error(f"单耗排名数据结构异常，缺少必要的列。现有列：{list(dept_unit_ranking.columns)}")
                display_unit_ranking_df = pd.DataFrame()
            
            try:
                st.dataframe(
                    display_unit_ranking_df.style.format({
                        '单耗': '¥{:.2f}',
                        '总成本': '¥{:.2f}万',
                        '产出': '{:.2f}万'
                    }),
                    use_container_width=True
                )
            except Exception:
                st.dataframe(display_unit_ranking_df, use_container_width=True)

def render_price_analysis_tab(tab, monthly_material_summary, df):
    """渲染单价分析标签页"""
    with tab:
        st.subheader("单价分析")
        
        if not monthly_material_summary.empty:
            # 按部门分组进行物料分析
            departments = monthly_material_summary['DEPT_NAME_STD'].unique()
            
            # 部门选择器
            selected_dept_analysis = st.selectbox(
                "选择部门进行物料单耗分析",
                options=['全部部门'] + list(departments),
                key='dept_material_analysis'
            )

            if selected_dept_analysis == '全部部门':
                # 显示所有部门的物料成本排名
                inv_ranking = (
                    monthly_material_summary.groupby(
                        ['INV', 'INV_PART_DESCRIPTION', 'DEPT_CODE_STD', 'DEPT_NAME_STD']
                    )
                    .agg({
                        'TOTAL_COST': 'sum',
                        'QUAN': 'sum',
                        'WPNL_SIZE': 'first',  # 产出
                        'UNIT_CONSUMPTION': 'mean'  # 平均单耗
                    })
                    .reset_index()
                    .sort_values('TOTAL_COST', ascending=False)
                )

                fig = px.bar(
                    inv_ranking.head(15),
                    x='INV_PART_DESCRIPTION',
                    y='TOTAL_COST',
                    color='DEPT_NAME_STD',
                    title="物料成本排名 (前15名)"
                )
                fig.update_xaxes(tickangle=45)
                st.plotly_chart(fig, use_container_width=True)

                # 物料详细表格
                st.subheader("物料汇总数据")
                required_inv_cols = [
                    'INV_PART_DESCRIPTION', 'DEPT_NAME_STD', 'TOTAL_COST', 'QUAN', 'WPNL_SIZE', 'UNIT_CONSUMPTION'
                ]
                if all(col in inv_ranking.columns for col in required_inv_cols):
                    display_inv_df = inv_ranking[required_inv_cols].copy()
                    display_inv_df['TOTAL_COST'] = display_inv_df['TOTAL_COST'] / 10000  # 转换为万元
                    display_inv_df['WPNL_SIZE'] = display_inv_df['WPNL_SIZE'] / 10000  # 转换为万
                    display_inv_df.columns = ['物料描述', '所属部门', '总成本', '总数量', '产出', '单耗']
                else:
                    st.error(
                        f"物料数据结构异常，缺少必要的列。现有列：{list(inv_ranking.columns)}"
                    )
                    display_inv_df = pd.DataFrame()

                try:
                    st.dataframe(
                        display_inv_df.style.format({
                            '总成本': '¥{:.2f}万',
                            '总数量': '{:,.0f}',
                            '产出': '{:.2f}万',
                            '单耗': '¥{:.2f}'
                        }),
                        use_container_width=True
                    )
                except Exception:
                    st.dataframe(display_inv_df, use_container_width=True)
            else:
                # 分析特定部门的物料单耗趋势
                dept_materials = monthly_material_summary[
                    monthly_material_summary['DEPT_NAME_STD'] == selected_dept_analysis
                ].copy()

                if not dept_materials.empty:
                    st.subheader(f"{selected_dept_analysis} - 物料单耗趋势分析")
                    
                    # 按物料和年月分析单耗趋势
                    material_trend = dept_materials.groupby(['INV', 'YEAR_MONTH']).agg({
                        'UNIT_CONSUMPTION': 'mean',
                        'TOTAL_COST': 'sum',
                        'QUAN': 'sum'
                    }).reset_index()
            
                    # 获取单耗最高的前10个物料
                    top_materials_unit = dept_materials.groupby('INV')['UNIT_CONSUMPTION'].mean().nlargest(10)
                    top_material_trend = material_trend[material_trend['INV'].isin(top_materials_unit.index)]
                    
                    if not top_material_trend.empty:
                        # 转换年月为字符串
                        top_material_trend['年月字符串'] = top_material_trend['YEAR_MONTH'].astype(str)
                        
                        # 绘制单耗趋势图
                        fig_unit_trend = px.line(
                            top_material_trend,
                            x='年月字符串',
                            y='UNIT_CONSUMPTION',
                            color='INV',
                            title=f'{selected_dept_analysis} - 前10大单耗物料趋势',
                            labels={'UNIT_CONSUMPTION': '单耗 (¥)', '年月字符串': '年月', 'INV': '物料'}
                        )
                        fig_unit_trend.update_layout(height=500)
                        st.plotly_chart(fig_unit_trend, use_container_width=True)
                    else:
                        st.info("该部门没有物料单耗趋势数据")
                else:
                    st.info("该部门没有物料数据")

def render_trend_analysis_tab(tab, df):
    """渲染月度趋势分析标签页"""
    with tab:
        st.subheader("单价变差分析")

        # 获取日期范围用于单价趋势查询
        if 'DATE' in df.columns:
            start_date = df['DATE'].min()
            end_date = df['DATE'].max()
        else:
            start_date = pd.to_datetime(df['YYYYMM'].min() + '01', format='%Y%m%d')
            end_date = pd.to_datetime(df['YYYYMM'].max() + '01', format='%Y%m%d') + pd.offsets.MonthEnd(0)

        # 加载所有料号的单价数据
        unit_price_df = load_unit_price_data(start_date, end_date)

        if not unit_price_df.empty:
            # 每日单价变差分析（跟前一天对比）
            st.subheader("每日单价变差前10名明细")
            
            # 计算每日单价变化（跟前一天对比）
            unit_price_daily = unit_price_df.sort_values(['INV', 'DATE'])
            unit_price_daily['前一天单价'] = unit_price_daily.groupby('INV')['UNIT_PRICE'].shift(1)
            unit_price_daily['日单价变差'] = unit_price_daily['UNIT_PRICE'] - unit_price_daily['前一天单价']
            unit_price_daily['日单价变差百分比'] = (unit_price_daily['日单价变差'] / unit_price_daily['前一天单价'] * 100).round(2)
            
            # 筛选最新日期的数据并按变差绝对值排序
            latest_date = unit_price_daily['DATE'].max()
            current_day_price_changes = unit_price_daily[
                (unit_price_daily['DATE'] == latest_date) & 
                (unit_price_daily['日单价变差百分比'].notna())
            ].copy()
            
            current_day_price_changes['单价变差绝对值'] = current_day_price_changes['日单价变差百分比'].abs()
            top_daily_price_changes = current_day_price_changes.nlargest(10, '单价变差绝对值')
            
            if not top_daily_price_changes.empty:
                display_daily_price_changes_df = top_daily_price_changes[['INV', 'DATE', '日单价变差百分比']].copy()
                display_daily_price_changes_df.columns = ['物料', '当前日期', '变差百分比(%)']
                
                try:
                    st.dataframe(
                        display_daily_price_changes_df.style.format({
                            '变差百分比(%)': '{:+.2f}%'
                        }),
                        use_container_width=True
                    )
                except Exception:
                    st.dataframe(display_daily_price_changes_df, use_container_width=True)
            else:
                st.info("没有足够的日度数据进行每日单价变差分析")
        else:
            st.warning("无法加载单价数据。")

def render_ai_analysis_tab(tab, df):
    """渲染AI分析标签页"""
    with tab:
        st.subheader("AI智能分析助手")
        st.caption("与AI对话，分析当前单耗数据")
        
        # LLM配置
        llm_model_name = "QWQ:32B" 
        llm_base_url = "http://172.18.80.88:8106/v1"
        
        # 记录数限制配置
        col1, col2 = st.columns([1, 3])
        with col1:
            max_rows = st.number_input(
                "分析最大数据量", 
                min_value=10, 
                max_value=1000,
                value=500,
                step=50,
                key="cost_llm_max_rows"
            )
        
        # 创建消息容器
        message_container = st.container()
        # 显示历史消息
        with message_container:
            for i, message in enumerate(st.session_state.cost_llm_messages):
                if message["role"] == "user":
                    # 用户消息显示在右侧
                    col1, col2 = st.columns([1, 3])
                    with col2:
                        st.markdown(f"""
                        <div style="background-color: #e3f2fd; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                            <strong>用户:</strong><br>
                            {message["content"]}
                        </div>
                        """, unsafe_allow_html=True)
                else:
                    # LLM消息显示在左侧
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.markdown(f"""
                        <div style="background-color: #f5f5f5; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                            <strong>AI助手:</strong><br>
                            {message["content"]}
                        </div>
                        """, unsafe_allow_html=True)
        
        # 使用streamlit的chat_input作为输入框
        user_input = st.chat_input("请分析当前单耗数据")
        
        # 处理用户输入
        if user_input:
            # 添加用户消息到会话状态
            st.session_state.cost_llm_messages.append({"role": "user", "content": user_input})
            
            # 显示用户消息
            with message_container:
                col1, col2 = st.columns([1, 3])
                with col2:
                    st.markdown(f"""
                    <div style="background-color: #e3f2fd; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                        <strong>用户:</strong><br>
                        {user_input}
                    </div>
                    """, unsafe_allow_html=True)
            
            # 显示助手消息
            with message_container:
                col1, col2 = st.columns([3, 1])
                with col1:
                    # 创建一个占位符用于流式输出
                    message_placeholder = st.empty()
                    message_placeholder.markdown("""
                    <div style="background-color: #f5f5f5; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                        <strong>AI助手:</strong><br>
                        正在分析中...
                    </div>
                    """, unsafe_allow_html=True)

                    try:
                        # 创建一个专用于LLM的数据副本
                        llm_df = df.copy()
                        
                        # 需要删除的列列表（分析生成的技术字段）
                        columns_to_remove = [
                            'IS_ANOMALY',
                            'Predicted', 
                            'Lower', 
                            'Upper', 
                            'Actual'
                        ]

                        # 从DataFrame中删除存在的列
                        for col in columns_to_remove:
                            if col in llm_df.columns:
                                llm_df.drop(col, axis=1, inplace=True)
                        
                        # 处理DATE列的格式转换和排序
                        if 'DATE' in llm_df.columns:
                            try:
                                # 确保DATE是datetime类型用于排序
                                llm_df['DATE'] = pd.to_datetime(llm_df['DATE'])
                                # 按时间从新到旧排序
                                llm_df = llm_df.sort_values('DATE', ascending=False)
                                # 限制为最新的max_rows条记录
                                llm_df_final = llm_df.head(max_rows)
                                # 将DATE转为字符串格式用于JSON序列化
                                llm_df_final['DATE'] = llm_df_final['DATE'].dt.strftime('%Y-%m-%d')
                            except Exception as date_err:
                                st.warning(f"处理DATE列时出错: {date_err}. 将使用未排序的数据。")
                                llm_df_final = llm_df.head(max_rows)
                        else:
                            # 按YYYYMM排序
                            try:
                                llm_df = llm_df.sort_values('YYYYMM', ascending=False)
                                llm_df_final = llm_df.head(max_rows)
                            except:
                                llm_df_final = llm_df.head(max_rows)
                        
                        # 将精简后的数据集转换为CSV格式
                        data_context = llm_df_final.to_csv(index=False)
                        
                        # 调用LLM
                        try:
                            system_prompt = """你是一个专业的成本分析师和数据分析助手。请基于提供的单耗分析数据和用户问题进行深入分析并回答。

                                数据字段说明：
                                - YYYYMM: 年月
                                - DEPT_NAME_STD: 部门名称
                                - INV: 物料品名
                                - INV_PART_DESCRIPTION: 物料描述
                                - QUAN: 数量
                                - STD_COST: 单价
                                - TOTAL_COST: 总成本
                                - WPNL_SIZE: 产出
                                - UNIT_CONSUMPTION: 单耗（总成本/产出）
                                - YEAR_MONTH: 年月期间
                                - DEPT_CODE_STD: 部门代码

                                请根据数据特点提供专业的成本分析建议。"""
                            
                            data_description = f"单耗分析数据：最新的{len(llm_df_final)}条记录\n数据时间范围：{llm_df_final.get('YYYYMM', pd.Series()).min() or '未知'} 至 {llm_df_final.get('YYYYMM', pd.Series()).max() or '未知'}\n\n"
                            
                            llm = ChatOpenAI(
                                model=llm_model_name,
                                openai_api_key="dummy-key",
                                openai_api_base=llm_base_url,
                                temperature=0,
                                streaming=True  # 启用流式输出
                            )
                            
                            enhanced_prompt = f"""
                                {data_description}以下是CSV格式的单耗分析数据：
                                ```csv
                                {data_context}
                                ```

                                用户问题: {user_input}

                                请根据以上单耗分析数据进行专业分析并回答。
                                """
                            
                            messages = [
                                SystemMessage(content=system_prompt),
                                HumanMessage(content=enhanced_prompt)
                            ]
                            
                            # 使用流式输出
                            llm_response_content = ""
                            for chunk in llm.stream(messages):
                                if chunk.content:
                                    llm_response_content += chunk.content
                                    message_placeholder.markdown(f"""
                                    <div style="background-color: #f5f5f5; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                                        <strong>AI助手:</strong><br>
                                        {llm_response_content}▌
                                    </div>
                                    """, unsafe_allow_html=True)
                            
                            # 移除光标并显示最终内容
                            message_placeholder.markdown(f"""
                            <div style="background-color: #f5f5f5; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                                <strong>AI助手:</strong><br>
                                {llm_response_content}
                            </div>
                            """, unsafe_allow_html=True)
                            
                        except Exception as llm_err:
                            llm_response_content = f"调用本地 LLM (模型: {llm_model_name} @ {llm_base_url}) 时出错: {llm_err}"
                            st.error(llm_response_content)
                            message_placeholder.markdown(f"""
                            <div style="background-color: #f5f5f5; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                                <strong>AI助手:</strong><br>
                                {llm_response_content}
                            </div>
                            """, unsafe_allow_html=True)

                    except Exception as e:
                        error_msg = f"处理数据时出错: {e}"
                        message_placeholder.markdown(f"""
                        <div style="background-color: #f5f5f5; padding: 10px; border-radius: 10px; margin: 5px 0; text-align: left;">
                            <strong>AI助手:</strong><br>
                            {error_msg}
                        </div>
                        """, unsafe_allow_html=True)
                        llm_response_content = error_msg
                    
                    # 添加助手消息到会话状态
                    st.session_state.cost_llm_messages.append({"role": "assistant", "content": llm_response_content})

def render_sidebar_filters():
    """渲染侧边栏筛选器"""
    with st.sidebar:
        # 数据库连接状态
        with st.status("连接数据库...", expanded=False) as status:
            try:
                with get_db_connection() as conn:
                    cursor = conn.cursor()
                    cursor.execute("SELECT 1 FROM DUAL")
                    cursor.fetchone()
                status.update(label="已连接数据库", state="complete", expanded=False)
            except Exception as e:
                status.update(label=f"数据库连接失败: {e}", state="error", expanded=True)
                st.stop()
        
        st.markdown("---")
        
        # 设置默认时间为最近三个月
        today = date.today()
        default_start = today - timedelta(days=90)  # 三个月前
        default_end = today
        
        # 日期输入控件
        col_start, col_end = st.columns(2)
        with col_start:
            start_date = st.date_input(
                "开始日期", 
                value=default_start
            )
        with col_end:
            end_date = st.date_input(
                "结束日期", 
                value=default_end
            )
         
        # 根据时间范围自动加载筛选选项
        with st.spinner("正在加载筛选选项..."):
            try:
                dept_options, inv_options = load_filter_options(start_date, end_date)
            except Exception as e:
                st.error(f"加载筛选选项失败: {e}")
                dept_options, inv_options = pd.DataFrame(), pd.DataFrame()
        
        # 初始化筛选选项
        selected_dept = "全部"
        selected_inv = "全部"
        
        if not dept_options.empty:
            # 部门筛选
            dept_list = ["全部"] + dept_options['DEPT_CODE_STD'].tolist()
            dept_format_func = lambda x: "全部" if x == "全部" else dept_options[dept_options['DEPT_CODE_STD']==x]['DEPT_NAME_STD'].iloc[0]
            
            selected_dept = st.selectbox(
                "部门选择",
                options=dept_list,
                format_func=dept_format_func,
                help=f"当前时间范围内共有 {len(dept_options)} 个部门"
            )
            
            # 根据选择的部门获取物料选项
            if selected_dept != "全部":
                try:
                    # 加载选定部门下的物料选项
                    with get_db_connection() as conn:
                        query = """
                        SELECT DISTINCT 品名 as INV
                        FROM VI_Wlcb_cost_ai
                        WHERE YYYYMM >= TO_CHAR(:start_date, 'YYYYMM')
                          AND YYYYMM <= TO_CHAR(:end_date, 'YYYYMM')
                          AND 部門名稱 = :dept
                        ORDER BY 品名
                        """
                        df_materials = pd.read_sql(query, conn, params={
                            'start_date': start_date,
                            'end_date': end_date,
                            'dept': selected_dept
                        })
                        if not df_materials.empty:
                            df_materials['INV_PART_DESCRIPTION'] = df_materials['INV']  # 使用品名作为显示名称
                            filtered_inv_options = df_materials
                        else:
                            filtered_inv_options = pd.DataFrame()
                except Exception as e:
                    st.warning(f"加载物料选项失败: {e}")
                    filtered_inv_options = inv_options
            else:
                filtered_inv_options = inv_options
            
            # 物料筛选
            if not filtered_inv_options.empty:
                inv_list = ["全部"] + filtered_inv_options['INV'].tolist()
                inv_format_func = lambda x: "全部" if x == "全部" else filtered_inv_options[filtered_inv_options['INV']==x]['INV_PART_DESCRIPTION'].iloc[0]
                
                material_help_text = f"当前部门下共有 {len(filtered_inv_options)} 个物料" if selected_dept != "全部" else f"所有部门下共有 {len(filtered_inv_options)} 个物料"
                selected_inv = st.selectbox(
                    "物料选择",
                    options=inv_list,
                    format_func=inv_format_func,
                    help=material_help_text
                )
            else:
                st.selectbox("物料选择", ["指定条件下无数据"], disabled=True)
        else:
            # 没有筛选选项时显示提示
            st.selectbox("部门选择", ["指定时间范围内无数据"], disabled=True)
            st.selectbox("物料选择", ["指定时间范围内无数据"], disabled=True)
            st.warning("指定时间范围内没有数据，请调整时间范围")
        
        # 按钮区域
        col_btn1, col_btn2 = st.columns(2)
        
        with col_btn1:
            # 加载数据按钮
            load_data = st.button("加载数据", use_container_width=True, type="primary", disabled=dept_options.empty)
        
        with col_btn2:
            # 重置筛选按钮
            if st.button("重置筛选", use_container_width=True):
                st.session_state.reset_filters = True
                st.rerun()
        
        return {
            'start_date': start_date,
            'end_date': end_date,
            'load_data': load_data,
            'selected_dept': selected_dept,
            'selected_inv': selected_inv,
            'dept_options': dept_options,
            'inv_options': inv_options
        }

def render_quantity_trend_analysis(df: pd.DataFrame):
    """渲染数量趋势分析"""
    if df.empty:
        st.warning("数量趋势数据为空")
        return
    
    # 数量趋势图
    fig = px.line(df, x='DATE', y='TOTAL_QUAN', 
                  title='总数量趋势', 
                  labels={'TOTAL_QUAN': '总数量', 'DATE': '日期'})
    fig.update_layout(height=400)
    st.plotly_chart(fig, use_container_width=True)
    
    # 显示统计信息
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("最大日数量", f"{df['TOTAL_QUAN'].max():,.0f}")
    with col2:
        st.metric("平均日数量", f"{df['TOTAL_QUAN'].mean():,.0f}")
    with col3:
        st.metric("总数量", f"{df['TOTAL_QUAN'].sum():,.0f}")

def render_unit_price_analysis(df: pd.DataFrame):
    """渲染单价分析"""
    if df.empty:
        st.warning("单价数据为空")
        return
    
    # 选择料号进行分析
    materials = df['INV'].unique()
    if len(materials) > 0:
        selected_material = st.selectbox("选择物料查看单价趋势", materials)
        
        material_data = df[df['INV'] == selected_material]
        if not material_data.empty:
            # 单价趋势图
            fig = px.line(material_data, x='DATE', y='UNIT_PRICE',
                         title=f'{selected_material} 单价趋势',
                         labels={'UNIT_PRICE': '单价 (¥)', 'DATE': '日期'})
            fig.update_layout(height=400)
            st.plotly_chart(fig, use_container_width=True)
            
            # 计算价格变化
            material_data = material_data.sort_values('DATE')
            if len(material_data) >= 2:
                first_price = material_data.iloc[0]['UNIT_PRICE']
                last_price = material_data.iloc[-1]['UNIT_PRICE']
                change_rate = ((last_price - first_price) / first_price * 100) if first_price != 0 else 0
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("起始单价", f"¥{first_price:.4f}")
                with col2:
                    st.metric("最新单价", f"¥{last_price:.4f}")
                with col3:
                    st.metric("变化率", f"{change_rate:+.2f}%")

def render_page_analysis(page_id: str, page_data: Dict[str, pd.DataFrame]):
    """根据页面类型渲染对应的分析结果"""
    
    if page_id == "cost_analysis":
        # 成本分析页面 - 使用综合数据或成本数据
        if 'comprehensive' in page_data and not page_data['comprehensive'].empty:
            render_analysis_results(page_data['comprehensive'])
        elif 'cost' in page_data and not page_data['cost'].empty:
            render_analysis_results(page_data['cost'])
        else:
            st.warning("暂无成本分析数据")
    
    elif page_id == "quantity_trend":
        # 数量趋势页面
        st.subheader("📈 数量趋势分析")
        if 'quantity_trend' in page_data and not page_data['quantity_trend'].empty:
            render_quantity_trend_analysis(page_data['quantity_trend'])
        if 'cost' in page_data and not page_data['cost'].empty:
            render_cost_quantity_analysis(page_data['cost'])
        if not any(page_data.values()):
            st.warning("暂无数量趋势数据")
    
    elif page_id == "price_analysis":
        # 单价分析页面
        st.subheader("💲 单价分析")
        if 'unit_price' in page_data and not page_data['unit_price'].empty:
            render_unit_price_analysis(page_data['unit_price'])
        if 'cost' in page_data and not page_data['cost'].empty:
            render_cost_price_analysis(page_data['cost'])
        if not any(page_data.values()):
            st.warning("暂无单价分析数据")
    
    elif page_id == "comprehensive":
        # 综合分析页面 - 显示所有数据类型
        st.subheader("🔍 综合分析")
        
        # 使用现有的render_analysis_results函数
        if 'comprehensive' in page_data and not page_data['comprehensive'].empty:
            render_analysis_results(page_data['comprehensive'])
        
        # 额外显示数量趋势
        if 'quantity_trend' in page_data and not page_data['quantity_trend'].empty:
            st.divider()
            st.subheader("数量趋势补充分析")
            render_quantity_trend_analysis(page_data['quantity_trend'])
        
        # 额外显示单价分析
        if 'unit_price' in page_data and not page_data['unit_price'].empty:
            st.divider()
            st.subheader("单价分析补充")
            render_unit_price_analysis(page_data['unit_price'])
        
        if not any(page_data.values()):
            st.warning("暂无综合分析数据")

def render_cost_quantity_analysis(df: pd.DataFrame):
    """基于成本数据渲染数量分析"""
    if df.empty or 'QUAN' not in df.columns:
        return
    
    st.subheader("基于成本数据的数量分析")
    
    # 按日期汇总数量
    if 'DATE' in df.columns:
        daily_quantity = df.groupby('DATE')['QUAN'].sum().reset_index()
        
        fig = px.line(daily_quantity, x='DATE', y='QUAN',
                     title='日度数量汇总趋势',
                     labels={'QUAN': '数量', 'DATE': '日期'})
        fig.update_layout(height=350)
        st.plotly_chart(fig, use_container_width=True)

def render_cost_price_analysis(df: pd.DataFrame):
    """基于成本数据渲染价格分析"""
    if df.empty or 'STD_COST' not in df.columns:
        return
    
    st.subheader("基于成本数据的价格分析")
    
    # 平均单价趋势
    if 'DATE' in df.columns and 'INV' in df.columns:
        # 选择前几个物料进行展示
        top_materials = df.groupby('INV')['TOTAL_COST'].sum().nlargest(5).index
        price_data = df[df['INV'].isin(top_materials)]
        
        if not price_data.empty:
            fig = px.line(price_data, x='DATE', y='STD_COST', color='INV',
                         title='主要物料单价趋势',
                         labels={'STD_COST': '单价 (¥)', 'DATE': '日期', 'INV': '物料'})
            fig.update_layout(height=350)
            st.plotly_chart(fig, use_container_width=True)

def main():
    """主函数 - 集成统一数据加载器和分页管理器"""
    
    # 渲染页面选择器
    current_page = page_manager.render_page_selector()
    page_config = page_manager.get_page_config(current_page)
    
    st.title(f"{page_config.title}")
    
    # 初始化session state
    session_key = f'{current_page}_data'
    filters_key = f'{current_page}_filters'
    messages_key = f'{current_page}_messages'
    
    if session_key not in st.session_state:
        st.session_state[session_key] = {}
    if filters_key not in st.session_state:
        st.session_state[filters_key] = {}
    if messages_key not in st.session_state:
        st.session_state[messages_key] = []
    
    # 处理重置筛选
    if st.session_state.get('reset_filters', False):
        # 清除当前页面相关的session state
        keys_to_remove = [key for key in st.session_state.keys() 
                         if key.startswith(current_page)]
        for key in keys_to_remove:
            del st.session_state[key]
        st.session_state.reset_filters = False
        st.session_state[session_key] = {}
        st.session_state[filters_key] = {}
        st.session_state[messages_key] = []
    
    # 渲染侧边栏筛选器
    filters = render_sidebar_filters()
    
    # 处理数据加载和分析
    if filters['load_data']:
        with st.spinner(f"正在加载{page_config.title}数据..."):
            
            # 构建筛选条件字典
            filter_dict = {}
            if filters.get('selected_dept') and filters['selected_dept'] != "全部":
                filter_dict['dept_filter'] = filters['selected_dept']
            if filters.get('selected_inv') and filters['selected_inv'] != "全部":
                filter_dict['inv_filter'] = filters['selected_inv']
            
            # 使用页面管理器加载数据
            page_data = page_manager.load_page_data(
                current_page, 
                filters['start_date'], 
                filters['end_date'],
                filter_dict if filter_dict else None
            )
            
            if page_data:
                # 检查是否有有效数据
                has_data = any(not df.empty for df in page_data.values())
                
                if has_data:
                    st.toast(f"成功加载{page_config.title}数据")
                    
                    # 存储数据和筛选条件到session state
                    st.session_state[session_key] = page_data
                    st.session_state[filters_key] = filters
                    
                else:
                    st.warning("筛选条件下没有匹配的数据，请调整筛选条件")
            else:
                st.warning("指定时间范围内没有数据")
    
    # 如果session state中有数据，显示分析结果
    if st.session_state[session_key]:
        render_page_analysis(current_page, st.session_state[session_key])

if __name__ == "__main__":
    main()